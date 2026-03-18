"""
ReportAnalyzer — تحليل تقرير Word سابق واستخراج الهيكل

يحلل ملف Word (.docx) ويستخرج:
1. Structure كل بند (ترتيب الفقرات والجداول والأشكال)
2. تعريف أعمدة كل جدول
3. أسلوب الكتابة (عينة نصية)

Usage:
    analyzer = ReportAnalyzer()
    result = analyzer.analyze_docx(file_path_or_bytes)

    # result = {
    #   'sections': [...],      # المحاور
    #   'style_sample': '...',  # عينة نصية
    #   'tables': [...],        # تعريفات الجداول
    #   'stats': {...},         # إحصائيات
    # }
"""

import io
import re
import logging
from typing import Optional, List, Dict, Any, Union

from docx import Document
from docx.table import Table as DocxTable

logger = logging.getLogger(__name__)


class ReportAnalyzer:
    """تحليل تقرير Word سابق واستخراج هيكله."""

    # Patterns to detect section/axis headings
    AXIS_PATTERNS = [
        re.compile(r'^المحور\s+(\d+|الأول|الثاني|الثالث|الرابع|الخامس|السادس|السابع|الثامن|التاسع|العاشر)\s*[:\-–—]\s*(.+)', re.UNICODE),
        re.compile(r'^محور\s+(\d+)\s*[:\-–—]\s*(.+)', re.UNICODE),
        re.compile(r'^(\d+)\s*[.\-–—]\s*(.+)', re.UNICODE),
    ]

    ITEM_PATTERNS = [
        re.compile(r'^(\d+\.\d+)\s*[:\-–—]?\s*(.+)', re.UNICODE),
        re.compile(r'^البند\s+(\d+\.\d+)\s*[:\-–—]?\s*(.+)', re.UNICODE),
    ]

    TABLE_TITLE_PATTERNS = [
        re.compile(r'جدول\s*\(?\s*(\d+[\-–]\d+)\s*\)?\s*[:\-–—]?\s*(.+)', re.UNICODE),
        re.compile(r'جدول\s+رقم\s*\(?\s*(\d+)\s*\)?\s*[:\-–—]?\s*(.+)', re.UNICODE),
    ]

    CHART_TITLE_PATTERNS = [
        re.compile(r'شكل\s*\(?\s*(\d+[\-–]\d+)\s*\)?\s*[:\-–—]?\s*(.+)', re.UNICODE),
        re.compile(r'شكل\s+رقم\s*\(?\s*(\d+)\s*\)?\s*[:\-–—]?\s*(.+)', re.UNICODE),
    ]

    def analyze_docx(self, source: Union[str, bytes, io.BytesIO]) -> Dict[str, Any]:
        """
        تحليل ملف Word.

        Args:
            source: مسار الملف أو bytes أو BytesIO

        Returns:
            dict مع sections, style_sample, tables, stats
        """
        try:
            if isinstance(source, str):
                doc = Document(source)
            elif isinstance(source, bytes):
                doc = Document(io.BytesIO(source))
            else:
                doc = Document(source)
        except Exception as e:
            logger.error(f'Failed to open document: {e}')
            return {'error': str(e), 'sections': [], 'style_sample': '', 'tables': [], 'stats': {}}

        sections = []
        all_tables = []
        style_samples = []
        current_axis = None
        current_item = None
        current_components = []

        paragraph_counter = 0
        table_counter = 0
        chart_counter = 0

        # Iterate through document body elements in order
        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for element in body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                if para_idx >= len(doc.paragraphs):
                    continue
                para = doc.paragraphs[para_idx]
                para_idx += 1

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ''
                is_heading = 'Heading' in style_name or style_name.startswith('heading')

                # Check if it's an axis heading
                axis_match = self._match_axis(text, is_heading)
                if axis_match:
                    # Save previous item
                    if current_item:
                        current_item['components'] = current_components
                        if current_axis:
                            current_axis['items'].append(current_item)
                        current_item = None
                        current_components = []

                    # Save previous axis
                    if current_axis:
                        sections.append(current_axis)

                    current_axis = {
                        'code': axis_match['code'],
                        'name': axis_match['name'],
                        'items': [],
                    }
                    continue

                # Check if it's an item heading
                item_match = self._match_item(text, is_heading)
                if item_match:
                    # Save previous item
                    if current_item:
                        current_item['components'] = current_components
                        if current_axis:
                            current_axis['items'].append(current_item)

                    paragraph_counter = 0
                    table_counter = 0
                    chart_counter = 0
                    current_components = []
                    current_item = {
                        'code': item_match['code'],
                        'name': item_match['name'],
                        'components': [],
                    }
                    continue

                # Check if it's a table title
                table_title_match = self._match_table_title(text)
                if table_title_match:
                    table_counter += 1
                    current_components.append({
                        'id': f't{table_counter}',
                        'type': 'table',
                        'title': table_title_match['title'],
                        'ref_number': table_title_match['number'],
                        'order': len(current_components) + 1,
                    })
                    continue

                # Check if it's a chart title
                chart_title_match = self._match_chart_title(text)
                if chart_title_match:
                    chart_counter += 1
                    current_components.append({
                        'id': f'c{chart_counter}',
                        'type': 'chart',
                        'title': chart_title_match['title'],
                        'ref_number': chart_title_match['number'],
                        'order': len(current_components) + 1,
                    })
                    continue

                # Regular paragraph
                if current_item and len(text) > 20:
                    paragraph_counter += 1
                    current_components.append({
                        'id': f'p{paragraph_counter}',
                        'type': 'paragraph',
                        'title': self._extract_paragraph_title(text),
                        'order': len(current_components) + 1,
                    })

                    # Collect style sample (first few substantial paragraphs)
                    if len(style_samples) < 3 and len(text) > 50:
                        style_samples.append(text)

            elif tag == 'tbl':
                if table_idx >= len(doc.tables):
                    continue
                docx_table = doc.tables[table_idx]
                table_idx += 1

                table_info = self._extract_table_info(docx_table)
                if table_info:
                    all_tables.append(table_info)

                    # If no table title was detected before this table, add one
                    if current_item:
                        last_comp = current_components[-1] if current_components else None
                        if not last_comp or last_comp['type'] != 'table':
                            table_counter += 1
                            current_components.append({
                                'id': f't{table_counter}',
                                'type': 'table',
                                'title': f'جدول {table_counter}',
                                'columns': table_info.get('headers', []),
                                'order': len(current_components) + 1,
                            })
                        elif last_comp and last_comp['type'] == 'table':
                            # Attach column info to the last table component
                            last_comp['columns'] = table_info.get('headers', [])

        # Save last item and axis
        if current_item:
            current_item['components'] = current_components
            if current_axis:
                current_axis['items'].append(current_item)

        if current_axis:
            sections.append(current_axis)

        # Build style sample
        style_sample = '\n\n'.join(style_samples)

        # Stats
        total_items = sum(len(s['items']) for s in sections)
        total_components = sum(
            len(item['components'])
            for s in sections
            for item in s['items']
        )
        total_paragraphs = sum(
            1
            for s in sections
            for item in s['items']
            for comp in item['components']
            if comp['type'] == 'paragraph'
        )
        total_tables_found = sum(
            1
            for s in sections
            for item in s['items']
            for comp in item['components']
            if comp['type'] == 'table'
        )

        return {
            'sections': sections,
            'style_sample': style_sample[:1000],
            'tables': all_tables,
            'stats': {
                'axes_count': len(sections),
                'items_count': total_items,
                'components_count': total_components,
                'paragraphs_count': total_paragraphs,
                'tables_count': total_tables_found,
                'extracted_tables': len(all_tables),
            },
        }

    # ==========================================
    # Pattern Matching
    # ==========================================

    def _match_axis(self, text: str, is_heading: bool) -> Optional[Dict]:
        """Match axis/section heading."""
        for pattern in self.AXIS_PATTERNS:
            m = pattern.match(text)
            if m:
                code = m.group(1)
                name = m.group(2).strip()
                # Convert Arabic ordinals to numbers
                ordinals = {
                    'الأول': '1', 'الثاني': '2', 'الثالث': '3', 'الرابع': '4',
                    'الخامس': '5', 'السادس': '6', 'السابع': '7', 'الثامن': '8',
                    'التاسع': '9', 'العاشر': '10',
                }
                code = ordinals.get(code, code)
                return {'code': code, 'name': name}

        # Heading style heuristic — short text in heading style
        if is_heading and len(text) < 100 and not any(c.isdigit() and '.' in text for c in text):
            return {'code': str(len(text) % 10 + 1), 'name': text}

        return None

    def _match_item(self, text: str, is_heading: bool) -> Optional[Dict]:
        """Match item heading (e.g., 1.9: عدد أعضاء هيئة التدريس)."""
        for pattern in self.ITEM_PATTERNS:
            m = pattern.match(text)
            if m:
                return {'code': m.group(1), 'name': m.group(2).strip()}
        return None

    def _match_table_title(self, text: str) -> Optional[Dict]:
        """Match table title (e.g., جدول (1-3): ...)."""
        for pattern in self.TABLE_TITLE_PATTERNS:
            m = pattern.search(text)
            if m:
                return {'number': m.group(1), 'title': m.group(2).strip()}
        return None

    def _match_chart_title(self, text: str) -> Optional[Dict]:
        """Match chart title (e.g., شكل (1-1): ...)."""
        for pattern in self.CHART_TITLE_PATTERNS:
            m = pattern.search(text)
            if m:
                return {'number': m.group(1), 'title': m.group(2).strip()}
        return None

    def _extract_paragraph_title(self, text: str) -> str:
        """Extract a short title from the first few words."""
        words = text.split()[:6]
        title = ' '.join(words)
        if len(title) > 40:
            title = title[:40] + '...'
        return title

    # ==========================================
    # Table Extraction
    # ==========================================

    def _extract_table_info(self, table: DocxTable) -> Optional[Dict]:
        """Extract headers and sample rows from a docx table."""
        try:
            rows = table.rows
            if len(rows) < 1:
                return None

            # First row = headers
            headers = [cell.text.strip() for cell in rows[0].cells]
            if not any(headers):
                return None

            # Data rows (up to 5 sample rows)
            data_rows = []
            for row in rows[1:6]:
                cells = [cell.text.strip() for cell in row.cells]
                data_rows.append(cells)

            # Detect column types
            column_types = []
            for col_idx, header in enumerate(headers):
                col_values = [
                    row[col_idx] for row in data_rows
                    if col_idx < len(row)
                ]
                col_type = self._detect_column_type(col_values)
                column_types.append({
                    'name': header,
                    'type': col_type,
                    'index': col_idx,
                })

            return {
                'headers': headers,
                'columns': column_types,
                'rows_count': len(rows) - 1,
                'sample_rows': data_rows,
            }
        except Exception as e:
            logger.warning(f'Failed to extract table: {e}')
            return None

    def _detect_column_type(self, values: List[str]) -> str:
        """Detect the type of data in a column."""
        if not values:
            return 'text'

        numeric_count = 0
        for v in values:
            v = v.replace(',', '').replace('٬', '').replace('%', '').strip()
            try:
                float(v)
                numeric_count += 1
            except ValueError:
                pass

        if numeric_count > len(values) * 0.5:
            # Check if percentage
            if any('%' in str(v) for v in values):
                return 'percentage'
            return 'number'

        return 'text'


def analyze_uploaded_report(file_bytes: bytes) -> Dict[str, Any]:
    """
    Convenience function to analyze an uploaded report.

    Args:
        file_bytes: The uploaded file content

    Returns:
        Analysis result dict
    """
    analyzer = ReportAnalyzer()
    return analyzer.analyze_docx(file_bytes)
