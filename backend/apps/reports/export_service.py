"""
Report Export Service
خدمة تصدير التقارير إلى Word/PDF

المميزات:
- تجميع المسودات (محاور + بنود)
- دمج المحتوى المولّد واليدوي
- تنسيق احترافي عربي
- دعم الجداول والصور
- توليد وإدراج الرسوم البيانية
"""

import os
import io
from typing import Optional, List, Dict, Any
from datetime import datetime
from django.conf import settings
from django.utils import timezone

# python-docx for Word generation
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from apps.data_collection.models import DataCollectionPeriod
from apps.reports.models import AxisDraft, ItemDraft, DraftAttachment

# Chart service
try:
    from apps.reports.chart_service import ChartService
    CHARTS_AVAILABLE = True
except ImportError:
    CHARTS_AVAILABLE = False


class ExportService:
    """خدمة تصدير التقارير"""
    
    def __init__(self, period: DataCollectionPeriod):
        self.period = period
        self.template = period.template
        self.doc = None
        
    def export_to_word(
        self,
        include_items: bool = True,
        include_tables: bool = True,
        include_charts: bool = True,
        include_attachments: bool = True,
        approved_only: bool = False,
    ) -> io.BytesIO:
        """
        تصدير التقرير إلى Word
        
        Args:
            include_items: تضمين مسودات البنود
            include_tables: تضمين الجداول
            include_charts: تضمين الرسوم البيانية
            include_attachments: تضمين المرفقات
            approved_only: فقط المحتوى المعتمد
        
        Returns:
            BytesIO containing the Word document
        """
        # Create document
        self.doc = Document()
        self._setup_document()
        
        # Add title page
        self._add_title_page()
        
        # Add table of contents placeholder
        self._add_toc_placeholder()
        
        # Get axis drafts
        axis_drafts = self.period.axis_drafts.select_related('axis').order_by('axis__order')
        
        if approved_only:
            axis_drafts = axis_drafts.filter(status='approved')
        
        # Add each axis
        for axis_draft in axis_drafts:
            self._add_axis_section(
                axis_draft,
                include_items=include_items,
                include_tables=include_tables,
                include_charts=include_charts,
                include_attachments=include_attachments,
                approved_only=approved_only,
            )
        
        # Save to BytesIO
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        
        return output
    
    def _setup_document(self):
        """إعداد المستند"""
        # Set RTL for Arabic
        sections = self.doc.sections
        for section in sections:
            section.page_width = Cm(21)  # A4
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
        
        # Setup styles
        self._setup_styles()
    
    def _setup_styles(self):
        """إعداد الأنماط"""
        styles = self.doc.styles
        
        # Title style (for main title)
        if 'Arabic Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Arabic Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style (for axes)
        if 'Arabic Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Arabic Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = 'Arial'
            h1_style.font.size = Pt(18)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(0, 51, 102)
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
        
        # Heading 2 style (for items)
        if 'Arabic Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Arabic Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = 'Arial'
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(0, 76, 153)
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(6)
        
        # Normal Arabic text
        if 'Arabic Normal' not in [s.name for s in styles]:
            normal_style = styles.add_style('Arabic Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Arial'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _set_rtl(self, paragraph):
        """Set RTL direction for paragraph"""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    def _add_title_page(self):
        """إضافة صفحة العنوان"""
        # University logo placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # TODO: Add logo image
        
        # Add spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Main title
        title = self.doc.add_paragraph()
        title.style = 'Arabic Title' if 'Arabic Title' in [s.name for s in self.doc.styles] else 'Title'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('التقرير السنوي')
        run.font.size = Pt(28)
        run.font.bold = True
        self._set_rtl(title)
        
        # Academic year
        year_p = self.doc.add_paragraph()
        year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = year_p.add_run(f'للعام الأكاديمي {self.period.academic_year}')
        run.font.size = Pt(20)
        run.font.bold = True
        self._set_rtl(year_p)
        
        # Template name
        if self.template:
            template_p = self.doc.add_paragraph()
            template_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = template_p.add_run(self.template.name)
            run.font.size = Pt(16)
            self._set_rtl(template_p)
        
        # Add spacing
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Date
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_date = timezone.now().strftime('%Y-%m-%d')
        run = date_p.add_run(f'تاريخ الإصدار: {current_date}')
        run.font.size = Pt(12)
        self._set_rtl(date_p)
        
        # Page break
        self.doc.add_page_break()
    
    def _add_toc_placeholder(self):
        """إضافة placeholder لجدول المحتويات"""
        toc_title = self.doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title.add_run('جدول المحتويات')
        run.font.size = Pt(18)
        run.font.bold = True
        self._set_rtl(toc_title)
        
        # Placeholder text
        placeholder = self.doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = placeholder.add_run('(سيتم توليده تلقائياً في Word)')
        run.font.size = Pt(10)
        run.font.italic = True
        self._set_rtl(placeholder)
        
        self.doc.add_page_break()
    
    def _add_axis_section(
        self,
        axis_draft: AxisDraft,
        include_items: bool = True,
        include_tables: bool = True,
        include_charts: bool = True,
        include_attachments: bool = True,
        approved_only: bool = False,
    ):
        """إضافة قسم محور"""
        axis = axis_draft.axis
        
        # Axis title
        title = self.doc.add_paragraph()
        title.style = 'Arabic Heading 1' if 'Arabic Heading 1' in [s.name for s in self.doc.styles] else 'Heading 1'
        run = title.add_run(f'المحور {axis.code}: {axis.name}')
        self._set_rtl(title)
        
        # Axis content (generated)
        if axis_draft.content:
            self._add_content_paragraphs(axis_draft.content)
        
        # Axis tables
        if include_tables and axis_draft.tables_data:
            for table_data in axis_draft.tables_data:
                self._add_table(table_data)
        
        # Axis attachments
        if include_attachments:
            for attachment in axis_draft.attachments.all():
                self._add_attachment(attachment)
        
        # Items
        if include_items:
            item_drafts = self.period.item_drafts.filter(
                item__axis=axis
            ).select_related('item').order_by('item__order')
            
            if approved_only:
                item_drafts = item_drafts.filter(status='approved')
            
            for item_draft in item_drafts:
                self._add_item_section(
                    item_draft,
                    include_tables=include_tables,
                    include_charts=include_charts,
                    include_attachments=include_attachments,
                )
        
        # Page break after each axis
        self.doc.add_page_break()
    
    def _add_item_section(
        self,
        item_draft: ItemDraft,
        include_tables: bool = True,
        include_charts: bool = True,
        include_attachments: bool = True,
    ):
        """إضافة قسم بند"""
        item = item_draft.item
        
        # Item title
        title = self.doc.add_paragraph()
        title.style = 'Arabic Heading 2' if 'Arabic Heading 2' in [s.name for s in self.doc.styles] else 'Heading 2'
        run = title.add_run(f'{item.code}. {item.name}')
        self._set_rtl(title)
        
        # Value info (if available)
        if item_draft.current_value is not None:
            value_p = self.doc.add_paragraph()
            value_text = f'القيمة الحالية: {item_draft.current_value}'
            if item.unit:
                value_text += f' {item.unit}'
            if item_draft.previous_value is not None:
                value_text += f' | القيمة السابقة: {item_draft.previous_value}'
            if item_draft.change_percentage is not None:
                sign = '+' if item_draft.change_percentage > 0 else ''
                value_text += f' | التغير: {sign}{item_draft.change_percentage}%'
            
            run = value_p.add_run(value_text)
            run.font.size = Pt(10)
            run.font.italic = True
            self._set_rtl(value_p)
        
        # Item content (generated)
        if item_draft.content:
            self._add_content_paragraphs(item_draft.content)
        
        # Table data (generated)
        if include_tables and item_draft.table_data:
            self._add_table_from_data(item_draft.table_data)
        
        # Chart (generated from data)
        if include_charts and CHARTS_AVAILABLE:
            self._add_item_chart(item_draft)
        
        # Manual content
        if item_draft.manual_content:
            self._add_manual_content(item_draft.manual_content)
        
        # Attachments
        if include_attachments:
            for attachment in item_draft.attachments.all():
                self._add_attachment(attachment)
    
    def _add_item_chart(self, item_draft: ItemDraft):
        """إضافة رسم بياني للبند"""
        try:
            chart_service = ChartService()
            chart_bytes = chart_service.generate_from_item_draft(item_draft)
            
            if chart_bytes:
                # إضافة عنوان الشكل
                chart_title = self.doc.add_paragraph()
                chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = chart_title.add_run(f'شكل ({item_draft.item.code}): {item_draft.item.name}')
                run.font.size = Pt(11)
                run.font.bold = True
                self._set_rtl(chart_title)
                
                # إضافة الرسم
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(chart_bytes), width=Inches(5.5))
                
                # مسافة بعد الرسم
                self.doc.add_paragraph()
        except Exception as e:
            # تجاهل الأخطاء في توليد الرسوم
            pass
    
    def _add_content_paragraphs(self, content: str):
        """إضافة فقرات المحتوى"""
        # Split by double newlines for paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = self.doc.add_paragraph()
                p.style = 'Arabic Normal' if 'Arabic Normal' in [s.name for s in self.doc.styles] else 'Normal'
                run = p.add_run(para_text.strip())
                self._set_rtl(p)
    
    def _add_table(self, table_config: dict):
        """إضافة جدول من تكوين"""
        title = table_config.get('title', '')
        headers = table_config.get('headers', [])
        rows = table_config.get('rows', [])
        
        if not rows:
            return
        
        # Table title
        if title:
            title_p = self.doc.add_paragraph()
            run = title_p.add_run(title)
            run.font.bold = True
            self._set_rtl(title_p)
        
        # Create table
        num_cols = len(headers) if headers else len(rows[0]) if rows else 1
        table = self.doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = str(header)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data if isinstance(row_data, list) else row_data.values()):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    cell.text = str(cell_data) if cell_data is not None else ''
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _add_table_from_data(self, table_data: list):
        """إضافة جدول من بيانات list"""
        if not table_data:
            return
        
        # Infer columns from first row
        first_row = table_data[0]
        if isinstance(first_row, dict):
            headers = list(first_row.keys())
            rows = [list(row.values()) for row in table_data]
        else:
            headers = []
            rows = table_data
        
        self._add_table({
            'headers': headers,
            'rows': rows,
        })
    
    def _add_manual_content(self, manual_content: list):
        """إضافة محتوى يدوي"""
        # Sort by order
        sorted_content = sorted(manual_content, key=lambda x: x.get('order', 0))
        
        for item in sorted_content:
            item_type = item.get('type')
            
            if item_type == 'text':
                self._add_content_paragraphs(item.get('content', ''))
            
            elif item_type == 'table':
                table_data = item.get('data', [])
                title = item.get('title', '')
                if table_data:
                    self._add_table({
                        'title': title,
                        'rows': table_data,
                    })
            
            elif item_type == 'image':
                # Handle image from attachment_id
                attachment_id = item.get('attachment_id')
                caption = item.get('caption', '')
                if attachment_id:
                    try:
                        attachment = DraftAttachment.objects.get(id=attachment_id)
                        self._add_attachment(attachment, caption_override=caption)
                    except DraftAttachment.DoesNotExist:
                        pass
    
    def _add_attachment(self, attachment: DraftAttachment, caption_override: str = None):
        """إضافة مرفق (صورة)"""
        if not attachment.file:
            return
        
        try:
            # Check if it's an image
            if attachment.file_type.startswith('image'):
                # Add image
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Get file path
                file_path = attachment.file.path
                if os.path.exists(file_path):
                    run = p.add_run()
                    run.add_picture(file_path, width=Inches(5))
                
                # Add caption
                caption = caption_override or attachment.caption
                if caption:
                    caption_p = self.doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption_p.add_run(caption)
                    run.font.size = Pt(10)
                    run.font.italic = True
                    self._set_rtl(caption_p)
                
                self.doc.add_paragraph()
        except Exception as e:
            # Log error but continue
            print(f"Error adding attachment {attachment.id}: {e}")


def export_period_to_word(
    period_id: int,
    include_items: bool = True,
    approved_only: bool = False,
) -> io.BytesIO:
    """
    Helper function to export a period to Word
    
    Args:
        period_id: ID of the DataCollectionPeriod
        include_items: Include item-level drafts
        approved_only: Only include approved content
    
    Returns:
        BytesIO containing the Word document
    """
    period = DataCollectionPeriod.objects.get(id=period_id)
    service = ExportService(period)
    return service.export_to_word(
        include_items=include_items,
        approved_only=approved_only,
    )
