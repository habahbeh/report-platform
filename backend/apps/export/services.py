"""
Export Service - Generate Word and PDF documents from reports.
"""

import os
import io
import json
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# WeasyPrint for PDF (optional)
try:
    from weasyprint import HTML, CSS
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False


def set_rtl_document(doc: Document):
    """Set document to RTL for Arabic."""
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)


def create_rtl_paragraph(doc: Document, text: str, style: str = None):
    """Create RTL paragraph."""
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set RTL
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    return p


def set_arabic_font(run, font_name: str = 'Arial', size: int = 12):
    """Set Arabic font for a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    r.get_or_add_rPr().append(rFonts)


def export_to_word(report) -> io.BytesIO:
    """
    Export report to Word document.
    
    Args:
        report: Report model instance
    
    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()
    
    # Set RTL
    set_rtl_document(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_arabic_font(run, 'Arial', 24)
    
    # Organization and period
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_name = report.organization.name if report.organization else 'جامعة البترا'
    run = subtitle.add_run(f'{org_name}\n{report.period_display}')
    set_arabic_font(run, 'Arial', 14)
    
    doc.add_paragraph()  # Spacer
    
    # Sections
    for section in report.sections.all().order_by('order'):
        # Section heading
        heading = doc.add_heading(section.title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in heading.runs:
            set_arabic_font(run, 'Arial', 18)
        
        # Section content
        if section.content:
            # Split by paragraphs
            paragraphs = section.content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = create_rtl_paragraph(doc, para_text.strip())
                    for run in p.runs:
                        set_arabic_font(run, 'Arial', 12)
        
        # Add images if any
        for image in section.images.all():
            if image.image and os.path.exists(image.image.path):
                doc.add_picture(image.image.path, width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if image.caption:
                    caption = create_rtl_paragraph(doc, image.caption)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        set_arabic_font(run, 'Arial', 10)
        
        doc.add_paragraph()  # Spacer between sections
    
    # Footer - Generated date
    footer_text = f'تم التوليد بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}'
    footer = create_rtl_paragraph(doc, footer_text)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_arabic_font(run, 'Arial', 10)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def export_to_pdf(report) -> Optional[io.BytesIO]:
    """
    Export report to PDF.
    
    Args:
        report: Report model instance
    
    Returns:
        BytesIO object containing the PDF, or None if WeasyPrint not available
    """
    if not HAS_WEASYPRINT:
        return None
    
    # Generate HTML
    html_content = generate_report_html(report)
    
    # CSS for RTL Arabic
    css = CSS(string='''
        @page {
            size: A4;
            margin: 2.5cm;
        }
        body {
            font-family: 'Arial', 'Tahoma', sans-serif;
            direction: rtl;
            text-align: right;
            line-height: 1.8;
            font-size: 12pt;
        }
        h1 {
            text-align: center;
            font-size: 24pt;
            margin-bottom: 0.5cm;
        }
        h2 {
            font-size: 18pt;
            margin-top: 1cm;
            border-bottom: 2px solid #333;
            padding-bottom: 0.3cm;
        }
        h3 {
            font-size: 14pt;
            margin-top: 0.5cm;
        }
        .subtitle {
            text-align: center;
            font-size: 14pt;
            margin-bottom: 1cm;
        }
        .section {
            margin-bottom: 1cm;
        }
        .image-container {
            text-align: center;
            margin: 1cm 0;
        }
        .image-container img {
            max-width: 80%;
            height: auto;
        }
        .caption {
            text-align: center;
            font-size: 10pt;
            color: #666;
            margin-top: 0.3cm;
        }
        .footer {
            text-align: center;
            font-size: 10pt;
            color: #999;
            margin-top: 2cm;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 1cm 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: right;
        }
        th {
            background-color: #f5f5f5;
        }
    ''')
    
    # Generate PDF
    buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(buffer, stylesheets=[css])
    buffer.seek(0)
    
    return buffer


def generate_report_html(report) -> str:
    """Generate HTML content for the report."""
    org_name = report.organization.name if report.organization else 'جامعة البترا'
    
    sections_html = ''
    for section in report.sections.all().order_by('order'):
        content = section.content.replace('\n', '<br>') if section.content else ''
        
        images_html = ''
        for image in section.images.all():
            if image.image:
                images_html += f'''
                <div class="image-container">
                    <img src="file://{image.image.path}" alt="{image.caption or 'صورة'}">
                    {f'<p class="caption">{image.caption}</p>' if image.caption else ''}
                </div>
                '''
        
        sections_html += f'''
        <div class="section">
            <h2>{section.title}</h2>
            <p>{content}</p>
            {images_html}
        </div>
        '''
    
    html = f'''
    <!DOCTYPE html>
    <html lang="ar" dir="rtl">
    <head>
        <meta charset="UTF-8">
        <title>{report.title}</title>
    </head>
    <body>
        <h1>{report.title}</h1>
        <p class="subtitle">{org_name}<br>{report.period_display}</p>
        
        {sections_html}
        
        <p class="footer">تم التوليد بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
    </body>
    </html>
    '''
    
    return html


def get_export_filename(report, format_type: str = 'docx') -> str:
    """Generate filename for export."""
    title_slug = report.title.replace(' ', '_')[:50]
    date_str = datetime.now().strftime('%Y%m%d')
    return f'{title_slug}_{date_str}.{format_type}'


# ============================================
# Project-based export (New System)
# ============================================

def export_project_to_word(project, generated_report=None, progress_callback=None) -> io.BytesIO:
    """
    Export project to Word document.
    Auto-detects V2 (ItemStructure) vs legacy (ItemDraft) system.
    """
    if _has_new_system_data(project):
        return export_project_v2_to_word(project, generated_report, progress_callback)
    from .report_template import generate_professional_report
    return generate_professional_report(project, generated_report)


def export_project_to_pdf(project, generated_report=None, progress_callback=None) -> Optional[io.BytesIO]:
    """
    Export project to PDF.
    Auto-detects V2 (ItemStructure) vs legacy (ItemDraft) system.
    """
    if _has_new_system_data(project):
        return export_project_v2_to_pdf(project, generated_report, progress_callback)

    if not HAS_WEASYPRINT:
        return None

    def update_progress(percent, step):
        if generated_report:
            generated_report.progress = percent
            generated_report.current_step = step
            generated_report.save(update_fields=['progress', 'current_step'])
        if progress_callback:
            progress_callback(percent, step)

    update_progress(10, 'تجهيز البيانات')

    html_content = generate_project_html(project, update_progress)

    update_progress(80, 'تحويل إلى PDF')

    css = CSS(string='''
        @page { size: A4; margin: 2.5cm; }
        body { font-family: 'Arial', 'Tahoma', sans-serif; direction: rtl;
               text-align: right; line-height: 1.8; font-size: 12pt; }
        h1 { text-align: center; font-size: 24pt; margin-bottom: 0.5cm; }
        h2 { font-size: 18pt; margin-top: 1cm; border-bottom: 2px solid #333; padding-bottom: 0.3cm; }
        h3 { font-size: 14pt; margin-top: 0.5cm; color: #444; }
        .subtitle { text-align: center; font-size: 14pt; margin-bottom: 1cm; }
        .value { font-size: 16pt; color: #0066cc; font-weight: bold; margin: 0.5cm 0; }
        .unit { font-size: 12pt; color: #666; }
        .footer { text-align: center; font-size: 10pt; color: #999; margin-top: 2cm; page-break-before: always; }
    ''')

    buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(buffer, stylesheets=[css])
    buffer.seek(0)

    update_progress(100, 'اكتمل')
    return buffer


def generate_pie_chart_css(data: list, title: str, chart_id: str = "1") -> str:
    """
    Generate CSS-based pie chart HTML.
    
    Args:
        data: List of dicts with 'label', 'value', 'color' (optional)
        title: Chart title
        chart_id: Unique identifier for the chart
    
    Returns:
        HTML string for the pie chart
    """
    # Default colors
    colors = ['#3b82f6', '#f59e0b', '#10b981', '#ef4444', '#8b5cf6', '#ec4899', '#06b6d4', '#84cc16']
    
    total = sum(item.get('value', 0) for item in data)
    if total == 0:
        return f'<div class="chart-container"><p>لا توجد بيانات</p></div>'
    
    # Calculate angles
    current_angle = 0
    gradient_parts = []
    legend_items = []
    
    for i, item in enumerate(data):
        value = item.get('value', 0)
        percentage = (value / total) * 100
        angle = (value / total) * 360
        color = item.get('color', colors[i % len(colors)])
        label = item.get('label', f'عنصر {i+1}')
        
        end_angle = current_angle + angle
        gradient_parts.append(f'{color} {current_angle}deg {end_angle}deg')
        current_angle = end_angle
        
        legend_items.append(f'''
            <div class="legend-item">
                <div class="legend-color" style="background: {color};"></div>
                <span>{label} {percentage:.1f}%</span>
            </div>
        ''')
    
    gradient = ', '.join(gradient_parts)
    legend_html = '\n'.join(legend_items)
    
    return f'''
    <div class="chart-container">
        <div class="chart-title">{title}</div>
        <div class="pie-chart" style="background: conic-gradient({gradient});"></div>
        <div class="legend">
            {legend_html}
        </div>
    </div>
    '''


def generate_bar_chart_css(data: list, title: str, chart_id: str = "1", max_height: int = 200) -> str:
    """
    Generate CSS-based bar chart HTML.
    
    Args:
        data: List of dicts with 'label', 'value', 'color' (optional)
        title: Chart title
        chart_id: Unique identifier
        max_height: Maximum bar height in pixels
    
    Returns:
        HTML string for the bar chart
    """
    colors = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6']
    
    if not data:
        return f'<div class="chart-container"><p>لا توجد بيانات</p></div>'
    
    max_value = max(item.get('value', 0) for item in data)
    if max_value == 0:
        max_value = 1
    
    bars_html = ''
    for i, item in enumerate(data):
        value = item.get('value', 0)
        height = int((value / max_value) * max_height)
        color = item.get('color', colors[i % len(colors)])
        label = item.get('label', '')
        
        bars_html += f'''
            <div class="bar-group">
                <div class="bar-value">{value:,}</div>
                <div class="bar" style="height: {height}px; background: {color};"></div>
                <div class="bar-label">{label}</div>
            </div>
        '''
    
    return f'''
    <div class="chart-container bar-chart-container">
        <div class="chart-title">{title}</div>
        <div class="bar-chart">
            {bars_html}
        </div>
    </div>
    '''


def generate_line_chart_css(data: list, title: str, chart_id: str = "1") -> str:
    """
    Generate CSS-based line chart HTML (using SVG).
    
    Args:
        data: List of dicts with 'label', 'value'
        title: Chart title
        chart_id: Unique identifier
    
    Returns:
        HTML string for the line chart
    """
    if not data or len(data) < 2:
        return f'<div class="chart-container"><p>لا توجد بيانات كافية</p></div>'
    
    width = 600
    height = 250
    padding = 50
    
    values = [item.get('value', 0) for item in data]
    max_val = max(values) if values else 1
    min_val = min(values) if values else 0
    range_val = max_val - min_val if max_val != min_val else 1
    
    # Calculate points
    points = []
    x_step = (width - 2 * padding) / (len(data) - 1)
    for i, item in enumerate(data):
        x = padding + i * x_step
        y = height - padding - ((item.get('value', 0) - min_val) / range_val) * (height - 2 * padding)
        points.append((x, y))
    
    # Create SVG path
    path_d = f"M {points[0][0]} {points[0][1]}"
    for x, y in points[1:]:
        path_d += f" L {x} {y}"
    
    # Create circles and labels
    circles_html = ''
    labels_html = ''
    for i, ((x, y), item) in enumerate(zip(points, data)):
        circles_html += f'<circle cx="{x}" cy="{y}" r="5" fill="#3b82f6"/>'
        labels_html += f'<text x="{x}" y="{height - 10}" text-anchor="middle" class="x-label">{item.get("label", "")}</text>'
    
    return f'''
    <div class="chart-container">
        <div class="chart-title">{title}</div>
        <svg viewBox="0 0 {width} {height}" class="line-chart-svg">
            <path d="{path_d}" fill="none" stroke="#3b82f6" stroke-width="3"/>
            {circles_html}
            {labels_html}
        </svg>
    </div>
    '''


def generate_table_html(data: dict, title: str = None) -> str:
    """
    Generate HTML table from data structure.
    
    Args:
        data: Dict with 'headers' (list) and 'rows' (list of lists)
        title: Optional table title
    
    Returns:
        HTML string for the table
    """
    if not data or 'rows' not in data:
        return ''
    
    headers = data.get('headers', [])
    rows = data.get('rows', [])
    
    title_html = f'<div class="table-title">{title}</div>' if title else ''
    
    header_html = ''
    if headers:
        header_cells = ''.join(f'<th>{h}</th>' for h in headers)
        header_html = f'<thead><tr>{header_cells}</tr></thead>'
    
    rows_html = ''
    for i, row in enumerate(rows):
        row_class = ''
        if isinstance(row, dict):
            row_class = row.get('class', '')
            cells = row.get('cells', [])
        else:
            cells = row
        
        cells_html = ''.join(f'<td>{cell}</td>' for cell in cells)
        rows_html += f'<tr class="{row_class}">{cells_html}</tr>'
    
    return f'''
    {title_html}
    <table>
        {header_html}
        <tbody>
            {rows_html}
        </tbody>
    </table>
    '''


def get_report_css() -> str:
    """Get the full CSS for professional Arabic reports."""
    return '''
        @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');
        
        * { box-sizing: border-box; }
        
        body {
            font-family: 'Cairo', 'Traditional Arabic', sans-serif;
            font-size: 16px;
            line-height: 1.8;
            max-width: 850px;
            margin: 40px auto;
            padding: 20px;
            direction: rtl;
            background: #fff;
            color: #333;
        }
        
        h1 {
            text-align: center;
            color: #1a365d;
            font-size: 28px;
            border-bottom: 3px solid #2563eb;
            padding-bottom: 15px;
            font-weight: 700;
        }
        
        h2 {
            color: #2563eb;
            font-size: 20px;
            margin-top: 30px;
            font-weight: 600;
        }
        
        h3 {
            color: #1a365d;
            font-size: 18px;
            margin-top: 20px;
            font-weight: 600;
        }
        
        p {
            text-align: justify;
            margin: 15px 0;
        }
        
        .subtitle {
            text-align: center;
            color: #666;
            font-size: 18px;
        }
        
        strong { color: #1a365d; }
        
        /* Tables */
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            font-size: 14px;
        }
        
        th, td {
            border: 1px solid #333;
            padding: 8px 10px;
            text-align: center;
        }
        
        th {
            background: #2563eb;
            color: white;
            font-weight: 600;
        }
        
        tr:nth-child(even) { background: #f0f4f8; }
        
        tr.total-row { 
            font-weight: bold; 
            background: #1a365d !important; 
            color: white;
        }
        
        tr.highlight-row {
            background: #dbeafe;
            font-weight: 600;
        }
        
        .table-title {
            font-weight: 700;
            margin: 30px 0 15px;
            font-size: 16px;
            color: #1a365d;
            text-align: center;
        }
        
        /* Charts */
        .chart-container {
            text-align: center;
            margin: 40px 0;
        }
        
        .chart-title {
            font-weight: 700;
            margin-bottom: 20px;
            font-size: 18px;
            color: #1a365d;
        }
        
        /* Pie Chart */
        .pie-chart {
            width: 280px;
            height: 280px;
            border-radius: 50%;
            margin: 0 auto 20px;
            position: relative;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }
        
        .pie-chart::before {
            content: '';
            position: absolute;
            width: 100px;
            height: 100px;
            background: white;
            border-radius: 50%;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
        }
        
        .legend {
            display: flex;
            justify-content: center;
            flex-wrap: wrap;
            gap: 20px;
            margin-top: 20px;
        }
        
        .legend-item {
            display: flex;
            align-items: center;
            gap: 10px;
            font-size: 14px;
            font-weight: 600;
        }
        
        .legend-color {
            width: 20px;
            height: 20px;
            border-radius: 4px;
        }
        
        /* Bar Chart */
        .bar-chart-container {
            padding: 20px;
        }
        
        .bar-chart {
            display: flex;
            justify-content: center;
            align-items: flex-end;
            gap: 30px;
            height: 250px;
            border-bottom: 2px solid #333;
            padding: 0 20px;
        }
        
        .bar-group {
            display: flex;
            flex-direction: column;
            align-items: center;
            gap: 5px;
        }
        
        .bar {
            width: 50px;
            border-radius: 4px 4px 0 0;
            transition: height 0.3s ease;
        }
        
        .bar-value {
            font-weight: 600;
            font-size: 14px;
            color: #1a365d;
        }
        
        .bar-label {
            font-size: 12px;
            color: #666;
            margin-top: 5px;
        }
        
        /* Line Chart */
        .line-chart-svg {
            width: 100%;
            max-width: 600px;
            height: auto;
        }
        
        .line-chart-svg .x-label {
            font-size: 12px;
            fill: #666;
        }
        
        /* Section */
        .section {
            margin-bottom: 40px;
            page-break-inside: avoid;
        }
        
        .item {
            margin: 20px 0;
            padding: 15px;
            background: #f8fafc;
            border-radius: 8px;
            border-right: 4px solid #2563eb;
        }
        
        .value {
            font-size: 18px;
            color: #0066cc;
            font-weight: bold;
            margin: 10px 0;
        }
        
        .unit {
            font-size: 14px;
            color: #666;
        }
        
        .footer {
            text-align: center;
            font-size: 12pt;
            color: #999;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }
        
        @media print {
            body { margin: 0; padding: 20px; }
            .pie-chart, tr.total-row {
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }
        }
    '''


def _parse_inline_tables(content: str) -> list:
    """
    Extract markdown tables from AI-generated content.
    Returns list of dicts: {'headers': [...], 'rows': [[...], ...]}
    """
    tables = []

    # Find markdown tables: lines starting with |
    # Pattern: | header | header | \n |---|---| \n | data | data |
    table_pattern = re.compile(
        r'(\|[^\n]+\|\n'           # header row
        r'\|[-:\s|]+\|\n'          # separator row
        r'(?:\|[^\n]+\|\n?)+)',    # data rows
        re.MULTILINE
    )

    for match in table_pattern.finditer(content):
        table_text = match.group(0)
        lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
        if len(lines) < 3:
            continue

        # Parse headers
        headers = [h.strip() for h in lines[0].split('|') if h.strip()]

        # Skip separator line, parse data rows
        rows = []
        for line in lines[2:]:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if len(cells) == len(headers):
                rows.append(cells)

        if headers and rows:
            tables.append({'headers': headers, 'rows': rows})

    return tables


def _parse_inline_charts(content: str) -> list:
    """
    Extract JSON chart configs from AI-generated content.
    Returns list of chart config dicts.
    """
    charts = []

    # Pattern 1: [CHART]\n{...json...} (with or without closing tag)
    chart_pattern = re.compile(
        r'\[CHART\]\s*\n?\s*(\{[^}]+\})',
        re.DOTALL
    )
    for match in chart_pattern.finditer(content):
        try:
            config = json.loads(match.group(1))
            charts.append(config)
        except json.JSONDecodeError:
            pass

    # Pattern 2: Standalone JSON with "type": "bar"|"pie"|"line"
    json_pattern = re.compile(
        r'\{"type":\s*"(bar|pie|line)"[^}]*\}',
        re.DOTALL
    )
    for match in json_pattern.finditer(content):
        try:
            config = json.loads(match.group(0))
            if config not in charts:
                charts.append(config)
        except json.JSONDecodeError:
            pass

    # Pattern 3: Multi-line JSON in ```json blocks
    json_block_pattern = re.compile(
        r'```json\s*\n\s*(\{[\s\S]*?\})\s*\n\s*```',
        re.DOTALL
    )
    for match in json_block_pattern.finditer(content):
        try:
            config = json.loads(match.group(1))
            if config.get('type') in ('bar', 'pie', 'line', 'gauge', 'info_card'):
                if config not in charts:
                    charts.append(config)
        except json.JSONDecodeError:
            pass

    # Pattern 4: Multi-line JSON objects with nested braces (charts with data objects)
    nested_json_pattern = re.compile(
        r'(\{\s*"type"\s*:\s*"(?:bar|pie|line)"[\s\S]*?\n\s*\})',
        re.MULTILINE
    )
    for match in nested_json_pattern.finditer(content):
        try:
            config = json.loads(match.group(1))
            if config not in charts:
                charts.append(config)
        except json.JSONDecodeError:
            pass

    return charts


def _render_chart_from_config(config: dict, item_code: str = '') -> str:
    """Render a chart config dict into HTML."""
    chart_type = config.get('type', 'bar')
    title = config.get('title', '')
    raw_data = config.get('data', {})

    # Skip non-chart types
    if chart_type in ('info_card',):
        return ''

    # Normalize data into list of {'label': ..., 'value': ...}
    chart_data = []

    if isinstance(raw_data, dict):
        labels = raw_data.get('labels', [])
        values = raw_data.get('values', raw_data.get('data', []))
        # Also check datasets format
        datasets = raw_data.get('datasets', [])
        if datasets and isinstance(datasets, list):
            values = datasets[0].get('data', datasets[0].get('values', []))

        colors_list = config.get('colors', [])
        default_colors = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#ec4899']

        for i, (label, value) in enumerate(zip(labels, values)):
            try:
                val = float(value)
            except (ValueError, TypeError):
                val = 0
            color = colors_list[i] if i < len(colors_list) else default_colors[i % len(default_colors)]
            chart_data.append({'label': str(label), 'value': val, 'color': color})
    elif isinstance(raw_data, list):
        for i, item in enumerate(raw_data):
            if isinstance(item, dict):
                chart_data.append(item)
            else:
                chart_data.append({'label': f'عنصر {i+1}', 'value': item})

    if not chart_data:
        return ''

    if chart_type == 'bar':
        return generate_bar_chart_css(chart_data, title, item_code)
    elif chart_type == 'pie':
        return generate_pie_chart_css(chart_data, title, item_code)
    elif chart_type == 'line':
        return generate_line_chart_css(chart_data, title, item_code)

    return ''


def _render_table_from_data(table_data: dict) -> str:
    """Render a parsed table dict into HTML."""
    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])

    if not rows:
        return ''

    header_html = ''
    if headers:
        header_cells = ''.join(f'<th>{h}</th>' for h in headers)
        header_html = f'<thead><tr>{header_cells}</tr></thead>'

    rows_html = ''
    for i, row in enumerate(rows):
        cells = ''.join(f'<td>{cell}</td>' for cell in row)
        row_class = ''
        # Detect total rows
        if any('مجموع' in str(c) or 'المجموع' in str(c) for c in row):
            row_class = 'total-row'
        rows_html += f'<tr class="{row_class}">{cells}</tr>'

    return f'''
    <div class="table-container">
        <table>
            {header_html}
            <tbody>{rows_html}</tbody>
        </table>
    </div>
    '''


def _format_content_to_html(content: str) -> str:
    """
    Convert AI-generated content to clean HTML.
    Extracts tables and charts, formats text into paragraphs.
    """
    if not content:
        return ''

    # 1. Extract tables and charts before cleaning
    tables = _parse_inline_tables(content)
    charts = _parse_inline_charts(content)

    # 2. Clean the text (remove markers, markdown tables, JSON blocks)
    text = content

    # Remove [TABLE], [CHART], [IMAGE] markers (with or without closing tags)
    text = re.sub(r'\[/?TABLE\]', '', text)
    text = re.sub(r'\[/?CHART\]', '', text)
    text = re.sub(r'\[/?IMAGE\]', '', text)
    text = re.sub(r'\[IMAGE\][\s\S]*?(?=\n\n|\[|\Z)', '', text)

    # Remove complete JSON blocks (```json...```)
    text = re.sub(r'```json[\s\S]*?```', '', text)

    # Remove standalone JSON objects with chart types
    text = re.sub(r'\{"type":\s*"(bar|pie|line|gauge|info_card)"[^}]*\}', '', text)
    text = re.sub(r'\{[\s\n]*"type":\s*"(bar|pie|line|gauge|info_card)"[\s\S]*?\n\s*\}', '', text)

    # Remove PARTIAL JSON fragments that AI leaves behind
    # e.g. ', "xAxis": "الفترة الزمنية", "yAxis": "عدد البرامج"}'
    # e.g. ', "colors": ["#6B7280", "#2563EB"]}'
    text = re.sub(r',\s*"[a-zA-Z_]+"\s*:\s*(?:"[^"]*"|[\d.]+|\[[^\]]*\]|\{[^}]*\})[^}\n]*\}?', '', text)

    # Remove any remaining JSON-like fragments: {"key": "value"...}
    text = re.sub(r'\{[^}]*"[a-zA-Z_]+"\s*:\s*[^}]+\}', '', text)

    # Remove orphaned JSON array fragments: ["...", "..."]
    text = re.sub(r'\["[^"]*"(?:\s*,\s*"[^"]*")*\]', '', text)

    # Remove markdown tables
    text = re.sub(
        r'\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)+',
        '', text, flags=re.MULTILINE
    )

    # Remove markdown formatting artifacts
    text = re.sub(r'^\*\*ملاحظة\*\*.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'---+', '', text)
    text = re.sub(r'^\s*📝\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*📊\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*📈\s*', '', text, flags=re.MULTILINE)

    # Convert headings
    text = re.sub(r'^### (.+)$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<h3 class="content-h3">\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.+)$', r'<h2 class="content-h2">\1</h2>', text, flags=re.MULTILINE)

    # Bold text
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)

    # Remove excess blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 3. Convert to paragraphs
    html_parts = []
    paragraphs = text.strip().split('\n\n')

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if p.startswith('<h'):
            html_parts.append(p)
        elif p.startswith('- ') or p.startswith('* '):
            items = p.split('\n')
            ul = '<ul>\n'
            for item in items:
                item_text = re.sub(r'^[-*]\s*', '', item.strip())
                if item_text:
                    ul += f'  <li>{item_text}</li>\n'
            ul += '</ul>'
            html_parts.append(ul)
        else:
            # Format numbers in parentheses as bold
            p = re.sub(r'\((\d[\d,\.]*)\)', r'<strong>(\1)</strong>', p)
            html_parts.append(f'<p>{p}</p>')

    result = '\n'.join(html_parts)

    # 4. Insert tables
    for table in tables:
        result += _render_table_from_data(table)

    # 5. Insert charts
    for chart in charts:
        result += _render_chart_from_config(chart)

    return result


def _generate_basic_html(project) -> str:
    """Generate a basic HTML report when no data models are available."""
    from datetime import datetime
    org_name = project.organization.name if project.organization else ''
    
    css = _get_enhanced_report_css()
    
    sections_html = ''
    if project.template:
        for axis in project.template.axes.all().order_by('order'):
            items_html = ''
            for item in axis.items.all().order_by('order'):
                items_html += f'''
                <article class="item" id="item-{item.code}">
                    <h3 class="item-title">{item.code}: {item.name}</h3>
                    <p class="placeholder">لم يتم إدخال بيانات بعد</p>
                </article>
'''
            sections_html += f'''
        <section class="axis" id="axis-{axis.code}">
            <h2 class="axis-title">المحور {axis.code}: {axis.name}</h2>
            {items_html}
        </section>
'''
    
    return f'''<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>{project.name}</title>
    <style>{css}</style>
</head>
<body>
    <header class="report-header">
        <h1>{project.name}</h1>
        <p class="subtitle">{org_name}</p>
        <p class="period">{project.period}</p>
    </header>
    <main class="report-content">
        {sections_html}
    </main>
    <footer class="report-footer">
        <p>تم توليد هذا التقرير بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
    </footer>
</body>
</html>'''


def generate_project_html(project, progress_callback=None) -> str:
    """Generate professional HTML content for the project report."""
    # Try V2 system first (ItemStructure-based)
    try:
        from apps.reports.models import ItemStructure
        if ItemStructure.objects.filter(project=project).exists():
            return generate_project_html_v2(project, progress_callback)
    except ImportError:
        pass

    # Legacy system - try to import old models
    try:
        from apps.reports.models import ItemDraft, AxisDraft
    except ImportError:
        # Old models don't exist - return basic HTML
        return _generate_basic_html(project)

    # Try to import DataCollectionPeriod (may not exist)
    try:
        from apps.data_collection.models import DataCollectionPeriod
        has_data_collection = True
    except ImportError:
        DataCollectionPeriod = None
        has_data_collection = False

    org_name = project.organization.name if project.organization else ''

    # Try project-based drafts first, fall back to period-based
    use_project_drafts = AxisDraft.objects.filter(project=project).exists()

    data_period = None
    if not use_project_drafts and has_data_collection and DataCollectionPeriod:
        data_period = DataCollectionPeriod.objects.filter(
            template=project.template,
            organization=project.organization,
            academic_year=project.period
        ).first()

    axes = project.template.axes.all().prefetch_related('items')
    total_axes = axes.count()

    sections_html = ''
    for i, axis in enumerate(axes):
        if progress_callback:
            progress = 10 + int((i / total_axes) * 60)
            progress_callback(progress, f'معالجة محور: {axis.name}')

        # Get axis draft for axis-level content
        axis_draft = None
        if use_project_drafts:
            axis_draft = AxisDraft.objects.filter(
                project=project,
                axis=axis
            ).first()
        elif data_period:
            axis_draft = AxisDraft.objects.filter(
                period=data_period,
                axis=axis
            ).first()

        items_html = ''
        for item in axis.items.all().order_by('order'):
            # Get item draft
            item_draft = None
            if use_project_drafts:
                item_draft = ItemDraft.objects.filter(
                    project=project,
                    item=item
                ).first()
            elif data_period:
                item_draft = ItemDraft.objects.filter(
                    period=data_period,
                    item=item
                ).first()

            if not item_draft:
                items_html += f'<article class="item" id="item-{item.code}"><h3 class="item-title">{item.code}: {item.name}</h3></article>'
                continue

            item_html = f'<article class="item" id="item-{item.code}">'
            item_html += f'<h3 class="item-title">{item.code}: {item.name}</h3>'

            # Render content: parse inline tables/charts from AI text
            if item_draft.content:
                item_html += f'<div class="item-text">{_format_content_to_html(item_draft.content)}</div>'

            # Also render structured chart_config if stored separately
            if item_draft.chart_config and isinstance(item_draft.chart_config, dict) and item_draft.chart_config.get('type'):
                item_html += _render_chart_from_config(item_draft.chart_config, item.code)

            # Also render structured table_data if stored separately
            if item_draft.table_data:
                td = item_draft.table_data
                if isinstance(td, list) and td:
                    if isinstance(td[0], dict):
                        headers = list(td[0].keys())
                        rows = [[row.get(h, '') for h in headers] for row in td]
                        item_html += _render_table_from_data({'headers': headers, 'rows': rows})
                    elif isinstance(td[0], list):
                        item_html += _render_table_from_data({'headers': td[0], 'rows': td[1:]})
                elif isinstance(td, dict) and td.get('rows'):
                    item_html += _render_table_from_data(td)

            # Manual content
            if item_draft.manual_content:
                item_html += f'<div class="manual-content">{item_draft.manual_content}</div>'

            item_html += '</article>'
            items_html += item_html

        # Build axis section
        axis_content = ''
        if axis_draft and axis_draft.content:
            axis_content = f'<div class="axis-text">{_format_content_to_html(axis_draft.content)}</div>'

        if items_html or axis_content:
            sections_html += f'''
        <section class="axis" id="axis-{axis.code}">
            <h2 class="axis-title">المحور {axis.code}: {axis.name}</h2>
            {axis_content}
            {items_html}
        </section>
            '''

    css = _get_enhanced_report_css()

    html = f'''<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.name}</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link href="https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap" rel="stylesheet">
    <style>{css}</style>
</head>
<body>
    <header class="report-header">
        <h1>{project.name}</h1>
        <p class="subtitle">{org_name}</p>
        <p class="period">{project.period}</p>
    </header>

    <main class="report-content">
        {sections_html}
    </main>

    <footer class="report-footer">
        <p>تم توليد هذا التقرير بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
    </footer>
</body>
</html>'''

    return html


def _get_enhanced_report_css() -> str:
    """Enhanced CSS for professional Arabic reports with charts and tables."""
    return '''
        @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap');

        * { box-sizing: border-box; margin: 0; padding: 0; }

        body {
            font-family: 'Cairo', 'Traditional Arabic', sans-serif;
            font-size: 16px;
            line-height: 2;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 30px;
            direction: rtl;
            background: #fff;
            color: #1a1a2e;
        }

        /* === Header === */
        .report-header {
            text-align: center;
            padding: 40px 20px;
            margin-bottom: 40px;
            border-bottom: 4px solid #1a365d;
            background: linear-gradient(135deg, #f8fafc 0%, #e8eef6 100%);
            border-radius: 12px;
        }

        .report-header h1 {
            color: #1a365d;
            font-size: 30px;
            font-weight: 900;
            margin-bottom: 10px;
            border: none;
            padding: 0;
        }

        .report-header .subtitle {
            color: #475569;
            font-size: 18px;
            font-weight: 600;
            margin: 5px 0;
        }

        .report-header .period {
            color: #2563eb;
            font-size: 16px;
            font-weight: 600;
            margin-top: 8px;
        }

        /* === Axes (Sections) === */
        .axis {
            margin-bottom: 50px;
            page-break-inside: avoid;
        }

        .axis-title {
            color: #fff;
            background: linear-gradient(135deg, #1a365d 0%, #2563eb 100%);
            font-size: 22px;
            font-weight: 700;
            padding: 15px 25px;
            border-radius: 10px;
            margin-bottom: 25px;
        }

        .axis-text {
            margin-bottom: 30px;
            padding: 0 10px;
        }

        .axis-text p {
            text-align: justify;
            margin: 15px 0;
            line-height: 2;
        }

        .axis-text h2.content-h2,
        .axis-text h3.content-h3,
        .axis-text h4 {
            color: #1a365d;
            margin-top: 25px;
            margin-bottom: 10px;
        }

        .axis-text h3.content-h3 { font-size: 18px; }
        .axis-text h4 { font-size: 16px; color: #2563eb; }

        /* === Items === */
        .item {
            margin: 25px 0;
            padding: 25px;
            background: #f8fafc;
            border-radius: 12px;
            border-right: 5px solid #2563eb;
            box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        }

        .item-title {
            color: #1a365d;
            font-size: 18px;
            font-weight: 700;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 2px solid #e2e8f0;
        }

        .item-text p {
            text-align: justify;
            margin: 12px 0;
            line-height: 2;
            color: #334155;
        }

        .item-text h3.content-h3,
        .item-text h4 {
            color: #1a365d;
            margin-top: 20px;
            margin-bottom: 8px;
        }

        .item-text h4 { font-size: 16px; color: #2563eb; }

        .item-text ul {
            margin: 10px 20px;
            padding: 0;
        }

        .item-text li {
            margin: 5px 0;
            line-height: 1.8;
        }

        strong { color: #1a365d; }

        /* === Tables === */
        .table-container {
            margin: 25px 0;
            overflow-x: auto;
        }

        table {
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
            border-radius: 8px;
            overflow: hidden;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        }

        th, td {
            border: 1px solid #d1d5db;
            padding: 10px 14px;
            text-align: center;
        }

        th {
            background: linear-gradient(135deg, #1a365d 0%, #2563eb 100%);
            color: white;
            font-weight: 700;
            font-size: 13px;
        }

        tr:nth-child(even) { background: #f1f5f9; }
        tr:hover { background: #e2e8f0; }

        tr.total-row {
            font-weight: bold;
            background: #1a365d !important;
            color: white;
        }

        .table-title {
            font-weight: 700;
            margin: 20px 0 10px;
            font-size: 15px;
            color: #1a365d;
            text-align: center;
        }

        /* === Charts === */
        .chart-container {
            text-align: center;
            margin: 30px auto;
            padding: 20px;
            background: #fff;
            border-radius: 12px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.06);
            max-width: 600px;
        }

        .chart-title {
            font-weight: 700;
            margin-bottom: 20px;
            font-size: 16px;
            color: #1a365d;
        }

        /* Pie Chart */
        .pie-chart {
            width: 260px;
            height: 260px;
            border-radius: 50%;
            margin: 0 auto 20px;
            position: relative;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }

        .pie-chart::before {
            content: '';
            position: absolute;
            width: 90px;
            height: 90px;
            background: white;
            border-radius: 50%;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
        }

        .legend {
            display: flex;
            justify-content: center;
            flex-wrap: wrap;
            gap: 15px;
            margin-top: 15px;
        }

        .legend-item {
            display: flex;
            align-items: center;
            gap: 8px;
            font-size: 13px;
            font-weight: 600;
        }

        .legend-color {
            width: 16px;
            height: 16px;
            border-radius: 4px;
            flex-shrink: 0;
        }

        /* Bar Chart */
        .bar-chart-container { padding: 15px; }

        .bar-chart {
            display: flex;
            justify-content: center;
            align-items: flex-end;
            gap: 25px;
            height: 220px;
            border-bottom: 2px solid #475569;
            padding: 0 15px;
        }

        .bar-group {
            display: flex;
            flex-direction: column;
            align-items: center;
            gap: 5px;
        }

        .bar {
            width: 55px;
            border-radius: 6px 6px 0 0;
            transition: height 0.3s ease;
        }

        .bar-value {
            font-weight: 700;
            font-size: 14px;
            color: #1a365d;
        }

        .bar-label {
            font-size: 12px;
            color: #64748b;
            margin-top: 8px;
            max-width: 80px;
            text-align: center;
        }

        /* Line Chart */
        .line-chart-svg {
            width: 100%;
            max-width: 550px;
            height: auto;
        }

        .x-label {
            font-size: 12px;
            fill: #64748b;
        }

        /* === Footer === */
        .report-footer {
            text-align: center;
            font-size: 13px;
            color: #94a3b8;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #e2e8f0;
        }

        /* === Print === */
        @media print {
            body { margin: 0; padding: 20px; max-width: 100%; }
            .report-header { background: #f8fafc !important; }
            .axis-title { background: #1a365d !important; }
            .item { box-shadow: none; border: 1px solid #e2e8f0; }
            .chart-container { box-shadow: none; border: 1px solid #e2e8f0; }
            th { background: #1a365d !important; }
            .pie-chart, tr.total-row, th {
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }
        }
    '''


def get_project_export_filename(project, format_type: str = 'docx') -> str:
    """Generate filename for project export."""
    title_slug = project.name.replace(' ', '_')[:50]
    date_str = datetime.now().strftime('%Y%m%d')
    return f'{title_slug}_{date_str}.{format_type}'


# ============================================
# V2 Export — SkeletonBuilder + ReferenceManager
# ============================================

def _has_new_system_data(project) -> bool:
    """Check if project uses the new ItemStructure system."""
    from apps.reports.models import ItemStructure
    return ItemStructure.objects.filter(project=project).exists()


def generate_project_html_v2(project, progress_callback=None) -> str:
    """
    Generate HTML using the new system:
    SkeletonBuilder (structure + data) + ReferenceManager (numbering) + GeneratedContent (AI text)
    """
    from .skeleton_builder import SkeletonBuilder
    from .reference_manager import ReferenceManager
    from apps.reports.models import ItemStructure, GeneratedContent

    def update_progress(percent, step):
        if progress_callback:
            progress_callback(percent, step)

    update_progress(5, 'تجهيز الهيكل')

    # 1. Build HTML skeleton (tables + charts + placeholders/generated text)
    builder = SkeletonBuilder(project)
    skeleton_html = builder.build_full_report(progress_callback=update_progress)

    update_progress(80, 'حل المراجع والترقيم')

    # 2. Resolve references {ref:t1} → جدول (1-3)
    ref_manager = ReferenceManager(project)
    ref_manager.build_registry()

    structures = ItemStructure.objects.filter(project=project)
    for structure in structures:
        contents = GeneratedContent.objects.filter(
            item_structure=structure,
            status__in=['generated', 'edited', 'approved'],
        )
        for gc in contents:
            if gc.content and '{ref:' in gc.content:
                resolved = ref_manager.resolve_references(gc.content, structure)
                # Replace in skeleton HTML
                skeleton_html = skeleton_html.replace(gc.content, resolved)
            if gc.manual_edit and '{ref:' in gc.manual_edit:
                resolved = ref_manager.resolve_references(gc.manual_edit, structure)
                skeleton_html = skeleton_html.replace(gc.manual_edit, resolved)

    update_progress(85, 'بناء الفهارس')

    # 3. Build Table of Contents
    toc_html = _build_toc_html(project)

    # 4. Build table/chart indexes
    table_index = ref_manager.get_table_index()
    chart_index = ref_manager.get_chart_index()
    indexes_html = _build_indexes_html(table_index, chart_index)

    # 5. Insert TOC after <main> and indexes before </main>
    if toc_html:
        skeleton_html = skeleton_html.replace(
            '<main>',
            f'<main>\n{toc_html}'
        ) if '<main>' in skeleton_html else skeleton_html
    if indexes_html:
        skeleton_html = skeleton_html.replace(
            '</main>',
            f'{indexes_html}\n    </main>'
        )

    # 5. Upgrade CSS — replace skeleton CSS with enhanced report CSS
    enhanced_css = _get_enhanced_report_css()
    skeleton_html = re.sub(
        r'<style>.*?</style>',
        f'<style>{enhanced_css}</style>',
        skeleton_html,
        flags=re.DOTALL,
    )

    # 6. Add Google Fonts link
    fonts_link = '<link rel="preconnect" href="https://fonts.googleapis.com">\n    <link href="https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap" rel="stylesheet">'
    skeleton_html = skeleton_html.replace(
        '<style>',
        f'{fonts_link}\n    <style>'
    )

    # 7. Add footer with generation date
    footer_html = f'''
    <footer class="report-footer">
        <p>تم توليد هذا التقرير بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
    </footer>'''
    skeleton_html = skeleton_html.replace('</main>', f'</main>\n{footer_html}')

    update_progress(90, 'اكتمل بناء HTML')

    # Log broken references
    broken = ref_manager.get_broken_references()
    if broken:
        import logging
        logger = logging.getLogger(__name__)
        for br in broken:
            logger.warning(
                f"مرجع مكسور في البند {br['item_code']}: {br['ref_id']}"
            )

    return skeleton_html


def _build_indexes_html(table_index: list, chart_index: list) -> str:
    """Build HTML for table and chart indexes (فهرس الجداول وفهرس الأشكال)."""
    if not table_index and not chart_index:
        return ''

    html = ''

    if table_index:
        html += '''
    <section class="index-section">
        <h2 class="index-title">فهرس الجداول</h2>
        <table class="index-table">
            <thead>
                <tr><th>الرقم</th><th>العنوان</th><th>البند</th></tr>
            </thead>
            <tbody>
'''
        for entry in table_index:
            html += f'                <tr><td>{entry["formal"]}</td><td>{entry["title"]}</td><td>{entry["item_code"]}: {entry["item_name"]}</td></tr>\n'
        html += '''            </tbody>
        </table>
    </section>
'''

    if chart_index:
        html += '''
    <section class="index-section">
        <h2 class="index-title">فهرس الأشكال</h2>
        <table class="index-table">
            <thead>
                <tr><th>الرقم</th><th>العنوان</th><th>البند</th></tr>
            </thead>
            <tbody>
'''
        for entry in chart_index:
            html += f'                <tr><td>{entry["formal"]}</td><td>{entry["title"]}</td><td>{entry["item_code"]}: {entry["item_name"]}</td></tr>\n'
        html += '''            </tbody>
        </table>
    </section>
'''

    return html


def _build_toc_html(project) -> str:
    """
    Build Table of Contents (فهرس المحتويات) — axes and items with anchor links.
    For HTML: uses anchor links. For PDF: WeasyPrint auto-generates page numbers.
    """
    from apps.templates_app.models import Axis

    axes = Axis.objects.filter(
        template=project.template
    ).prefetch_related('items').order_by('order')

    if not axes.exists():
        return ''

    html = '''
    <section class="toc-section" style="page-break-after: always;">
        <h2 class="index-title" style="text-align: center; margin-bottom: 30px;">فهرس المحتويات</h2>
        <div class="toc-content">
'''

    for axis in axes:
        html += f'''
            <div class="toc-axis" style="margin-bottom: 15px;">
                <div style="font-weight: 700; font-size: 1.05em; color: #1a365d; padding: 8px 0; border-bottom: 2px solid #e2e8f0;">
                    <a href="#axis-{axis.code}" style="text-decoration: none; color: inherit;">
                        المحور {axis.code}: {axis.name}
                    </a>
                </div>
                <div style="padding-right: 20px;">
'''
        for item in axis.items.order_by('order'):
            html += f'''
                    <div style="display: flex; justify-content: space-between; padding: 4px 0; border-bottom: 1px dotted #e2e8f0;">
                        <a href="#item-{item.code}" style="text-decoration: none; color: #4a5568; font-size: 0.95em;">
                            {item.code}: {item.name}
                        </a>
                    </div>
'''
        html += '''
                </div>
            </div>
'''

    html += '''
        </div>
    </section>
'''
    return html


def export_project_v2_to_word(project, generated_report=None, progress_callback=None) -> io.BytesIO:
    """
    Export project to Word using the new V2 system.
    Generates HTML first, then converts to Word via python-docx.
    """
    from bs4 import BeautifulSoup

    def update_progress(percent, step):
        if generated_report:
            generated_report.progress = percent
            generated_report.current_step = step
            generated_report.save(update_fields=['progress', 'current_step'])
        if progress_callback:
            progress_callback(percent, step)

    # Generate HTML
    html_content = generate_project_html_v2(project, update_progress)

    update_progress(90, 'تحويل إلى Word')

    # Parse HTML and build Word document
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    set_rtl_document(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    header = soup.find('header')
    if header:
        h1 = header.find('h1')
        if h1:
            title = doc.add_heading(h1.get_text(), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                set_arabic_font(run, 'Arial', 24)

        subtitle = header.find('p', class_='org-name') or header.find('p', class_='subtitle')
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle.get_text())
            set_arabic_font(run, 'Arial', 14)

    doc.add_paragraph()  # Spacer

    # Process sections (axes)
    for axis_section in soup.find_all('section', class_='axis'):
        axis_title = axis_section.find(class_='axis-title')
        if axis_title:
            heading = doc.add_heading(axis_title.get_text(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in heading.runs:
                set_arabic_font(run, 'Arial', 18)

        # Process items
        for item_article in axis_section.find_all('article', class_='item'):
            item_title = item_article.find(class_='item-title')
            if item_title:
                heading = doc.add_heading(item_title.get_text(), level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in heading.runs:
                    set_arabic_font(run, 'Arial', 16)

            # Sub-headings
            for h4 in item_article.find_all('h4', class_='sub-heading'):
                heading = doc.add_heading(h4.get_text(), level=3)
                heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in heading.runs:
                    set_arabic_font(run, 'Arial', 14)

            # Paragraphs (generated text)
            for para_div in item_article.find_all('div', class_='paragraph'):
                if 'placeholder' in para_div.get('class', []):
                    continue  # Skip unfilled placeholders
                for p_tag in para_div.find_all('p'):
                    text = p_tag.get_text()
                    if text.strip():
                        p = create_rtl_paragraph(doc, text.strip())
                        for run in p.runs:
                            set_arabic_font(run, 'Arial', 12)

            # Tables
            for table_div in item_article.find_all('div', class_='table-container'):
                if 'placeholder' in table_div.get('class', []):
                    continue
                table_title = table_div.find(class_='table-title')
                if table_title:
                    p = create_rtl_paragraph(doc, table_title.get_text())
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_arabic_font(run, 'Arial', 12)
                        run.bold = True

                html_table = table_div.find('table')
                if html_table:
                    rows = html_table.find_all('tr')
                    if rows:
                        # Count columns
                        first_row_cells = rows[0].find_all(['th', 'td'])
                        col_count = len(first_row_cells)
                        if col_count > 0:
                            word_table = doc.add_table(rows=0, cols=col_count)
                            word_table.style = 'Table Grid'

                            for row_el in rows:
                                cells = row_el.find_all(['th', 'td'])
                                row = word_table.add_row()
                                for i, cell in enumerate(cells):
                                    if i < col_count:
                                        row.cells[i].text = cell.get_text().strip()
                                        for paragraph in row.cells[i].paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                set_arabic_font(run, 'Arial', 11)
                                                if row_el.find('th'):
                                                    run.bold = True

        doc.add_paragraph()  # Spacer between axes

    # Footer
    footer_text = f'تم التوليد بتاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}'
    footer = create_rtl_paragraph(doc, footer_text)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_arabic_font(run, 'Arial', 10)

    update_progress(100, 'اكتمل')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_project_v2_to_pdf(project, generated_report=None, progress_callback=None) -> Optional[io.BytesIO]:
    """Export project to PDF using the new V2 system."""
    if not HAS_WEASYPRINT:
        return None

    def update_progress(percent, step):
        if generated_report:
            generated_report.progress = percent
            generated_report.current_step = step
            generated_report.save(update_fields=['progress', 'current_step'])
        if progress_callback:
            progress_callback(percent, step)

    html_content = generate_project_html_v2(project, update_progress)

    update_progress(90, 'تحويل إلى PDF')

    buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(buffer)
    buffer.seek(0)

    update_progress(100, 'اكتمل')
    return buffer
