"""
Item Generator - Production-ready report item generation.

Approach: Data First → Charts → Text → Export
"""

import os
import io
import json
from pathlib import Path
from typing import Optional, Dict, List, Any
from dataclasses import dataclass

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .chart_generator import generate_chart_image


@dataclass
class ItemComponent:
    """Represents a single component in an item structure."""
    id: str
    type: str  # 'paragraph', 'chart', 'table'
    title: str = ""
    content: str = ""
    data: Any = None
    image_path: str = ""
    status: str = "pending"  # pending, generated, approved


@dataclass
class GeneratedItem:
    """Represents a fully generated report item."""
    item_id: str
    item_code: str
    title: str
    components: List[ItemComponent]
    html: str = ""
    word_buffer: Optional[io.BytesIO] = None


class ItemGenerator:
    """
    Generates report items following the data-first approach.
    
    Flow:
    1. Load item structure from DB
    2. Fetch data for tables/charts
    3. Generate chart images
    4. Generate text using AI (sees all data)
    5. Build HTML skeleton
    6. Export to Word/PDF
    """
    
    def __init__(self, project_id: str, output_dir: str = None):
        self.project_id = project_id
        self.output_dir = Path(output_dir) if output_dir else Path("./output")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_item(self, item_structure) -> GeneratedItem:
        """
        Generate a complete item with all components.
        
        Args:
            item_structure: ItemStructure model instance
        
        Returns:
            GeneratedItem with HTML and Word document
        """
        from apps.templates_app.models import Item
        
        item = item_structure.item
        components = []
        
        # Parse structure
        structure = item_structure.structure or []
        
        for comp_def in structure:
            comp = self._process_component(comp_def, item_structure)
            components.append(comp)
        
        # Build result
        result = GeneratedItem(
            item_id=str(item.id),
            item_code=item.code,
            title=item.name,
            components=components
        )
        
        # Generate HTML
        result.html = self._build_html(result)
        
        # Generate Word
        result.word_buffer = self._build_word(result)
        
        return result
    
    def _process_component(self, comp_def: dict, item_structure) -> ItemComponent:
        """Process a single component definition."""
        comp_id = comp_def.get('id', '')
        comp_type = comp_def.get('type', 'paragraph')
        title = comp_def.get('title', '')
        
        component = ItemComponent(
            id=comp_id,
            type=comp_type,
            title=title
        )
        
        if comp_type == 'paragraph':
            component = self._process_paragraph(component, comp_def, item_structure)
        elif comp_type == 'chart':
            component = self._process_chart(component, comp_def, item_structure)
        elif comp_type == 'table':
            component = self._process_table(component, comp_def, item_structure)
        
        return component
    
    def _process_paragraph(self, component: ItemComponent, comp_def: dict, 
                           item_structure) -> ItemComponent:
        """Process paragraph component - fetch or generate text."""
        from apps.reports.models import ItemDraft
        
        # Try to get from ItemDraft
        try:
            draft = ItemDraft.objects.filter(
                item_structure=item_structure,
                status='approved'
            ).first()
            
            if draft and draft.content:
                component.content = draft.content
                component.status = 'approved'
        except:
            pass
        
        # If no content, mark as pending for AI generation
        if not component.content:
            component.status = 'pending'
        
        return component
    
    def _process_chart(self, component: ItemComponent, comp_def: dict,
                       item_structure) -> ItemComponent:
        """Process chart component - generate image."""
        from apps.templates_app.models import TableDefinition, TableData
        
        chart_config = comp_def.get('config', {})
        table_def_id = comp_def.get('table_def_id')
        
        # Get data from TableData
        labels = []
        values = []
        
        if table_def_id:
            try:
                table_data = TableData.objects.filter(
                    table_definition_id=table_def_id,
                    project_id=self.project_id
                ).first()
                
                if table_data and table_data.data:
                    data = table_data.data
                    if isinstance(data, list) and data:
                        # Extract labels and values
                        for row in data:
                            if isinstance(row, dict):
                                # Find label and value columns
                                for key, val in row.items():
                                    if 'year' in key.lower() or 'سنة' in key:
                                        labels.append(str(val))
                                    elif isinstance(val, (int, float)):
                                        values.append(val)
            except Exception as e:
                print(f"Error fetching chart data: {e}")
        
        # Generate chart image
        if labels and values:
            chart_type = chart_config.get('chart_type', 'bar')
            
            img_config = {
                'type': chart_type,
                'title': component.title,
                'data': {
                    'labels': labels,
                    'datasets': [{'label': '', 'values': values}]
                }
            }
            
            img_buffer = generate_chart_image(img_config)
            
            if img_buffer:
                # Save image
                img_path = self.output_dir / f"{component.id}.png"
                with open(img_path, 'wb') as f:
                    f.write(img_buffer.read())
                
                component.image_path = str(img_path)
                component.status = 'generated'
                component.data = {'labels': labels, 'values': values}
        
        return component
    
    def _process_table(self, component: ItemComponent, comp_def: dict,
                       item_structure) -> ItemComponent:
        """Process table component - fetch data."""
        from apps.templates_app.models import TableDefinition, TableData
        
        table_def_id = comp_def.get('table_def_id')
        
        if table_def_id:
            try:
                table_data = TableData.objects.filter(
                    table_definition_id=table_def_id,
                    project_id=self.project_id
                ).first()
                
                if table_data and table_data.data:
                    component.data = table_data.data
                    component.status = 'generated'
            except Exception as e:
                print(f"Error fetching table data: {e}")
        
        return component
    
    def _build_html(self, item: GeneratedItem) -> str:
        """Build HTML skeleton from components."""
        html_parts = [
            '<!DOCTYPE html>',
            '<html dir="rtl" lang="ar">',
            '<head>',
            '  <meta charset="UTF-8">',
            '  <title>{}</title>'.format(item.title),
            '  <style>',
            '    body { font-family: Arial, sans-serif; direction: rtl; padding: 20px; }',
            '    .item { max-width: 900px; margin: 0 auto; }',
            '    .paragraph { margin: 15px 0; line-height: 1.8; }',
            '    .chart-container { text-align: center; margin: 20px 0; }',
            '    .chart-container img { max-width: 100%; }',
            '    .table-container { margin: 20px 0; overflow-x: auto; }',
            '    table { width: 100%; border-collapse: collapse; }',
            '    th, td { border: 1px solid #ddd; padding: 8px; text-align: center; }',
            '    th { background: #f5f5f5; font-weight: bold; }',
            '    .placeholder { background: #fff3cd; padding: 10px; border-radius: 5px; }',
            '  </style>',
            '</head>',
            '<body>',
            '  <article class="item">',
            '    <h1>{}</h1>'.format(item.title),
        ]
        
        for comp in item.components:
            html_parts.append(self._component_to_html(comp))
        
        html_parts.extend([
            '  </article>',
            '</body>',
            '</html>'
        ])
        
        return '\n'.join(html_parts)
    
    def _component_to_html(self, comp: ItemComponent) -> str:
        """Convert component to HTML."""
        if comp.type == 'paragraph':
            if comp.content:
                return f'    <div class="paragraph" data-component="{comp.id}"><p>{comp.content}</p></div>'
            else:
                return f'    <div class="paragraph placeholder" data-component="{comp.id}"><span>في انتظار التوليد</span></div>'
        
        elif comp.type == 'chart':
            if comp.image_path:
                return f'''    <div class="chart-container" data-component="{comp.id}">
      <img src="{comp.image_path}" alt="{comp.title}">
      <p><em>{comp.title}</em></p>
    </div>'''
            else:
                return f'    <div class="chart-container placeholder" data-component="{comp.id}"><span>{comp.title} - في انتظار البيانات</span></div>'
        
        elif comp.type == 'table':
            if comp.data:
                return self._table_to_html(comp)
            else:
                return f'    <div class="table-container placeholder" data-component="{comp.id}"><span>{comp.title} - في انتظار البيانات</span></div>'
        
        return ''
    
    def _table_to_html(self, comp: ItemComponent) -> str:
        """Convert table component to HTML."""
        data = comp.data
        if not data or not isinstance(data, list):
            return ''
        
        html = [f'    <div class="table-container" data-component="{comp.id}">']
        
        if comp.title:
            html.append(f'      <p class="table-title"><strong>{comp.title}</strong></p>')
        
        html.append('      <table>')
        
        # Headers from first row keys
        if data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            html.append('        <thead><tr>')
            for h in headers:
                html.append(f'          <th>{h}</th>')
            html.append('        </tr></thead>')
            
            # Rows (limit to 50 for preview)
            html.append('        <tbody>')
            for row in data[:50]:
                html.append('          <tr>')
                for h in headers:
                    val = row.get(h, '')
                    html.append(f'            <td>{val}</td>')
                html.append('          </tr>')
            html.append('        </tbody>')
            
            if len(data) > 50:
                html.append(f'        <tfoot><tr><td colspan="{len(headers)}">... و {len(data) - 50} صف إضافي</td></tr></tfoot>')
        
        html.append('      </table>')
        html.append('    </div>')
        
        return '\n'.join(html)
    
    def _build_word(self, item: GeneratedItem) -> io.BytesIO:
        """Build Word document from components."""
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Title
        title = doc.add_heading(item.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_arabic_font(title.runs[0] if title.runs else None, 18)
        
        doc.add_paragraph()
        
        # Components
        for comp in item.components:
            self._add_component_to_word(doc, comp)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_component_to_word(self, doc: Document, comp: ItemComponent):
        """Add component to Word document."""
        if comp.type == 'paragraph' and comp.content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._set_rtl(p)
            run = p.add_run(comp.content)
            self._set_arabic_font(run, 12)
        
        elif comp.type == 'chart' and comp.image_path and os.path.exists(comp.image_path):
            doc.add_picture(comp.image_path, width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if comp.title:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(comp.title)
                self._set_arabic_font(run, 10)
                run.italic = True
            
            doc.add_paragraph()
        
        elif comp.type == 'table' and comp.data:
            self._add_table_to_word(doc, comp)
    
    def _add_table_to_word(self, doc: Document, comp: ItemComponent):
        """Add table to Word document."""
        data = comp.data
        if not data or not isinstance(data, list) or not isinstance(data[0], dict):
            return
        
        # Title
        if comp.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(comp.title)
            self._set_arabic_font(run, 12)
            run.bold = True
        
        headers = list(data[0].keys())
        max_rows = min(len(data), 30)  # Limit rows
        
        table = doc.add_table(rows=max_rows + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Headers
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = str(h)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    self._set_arabic_font(run, 9)
                    run.bold = True
        
        # Data rows
        for i, row in enumerate(data[:max_rows]):
            for j, h in enumerate(headers):
                cell = table.cell(i + 1, j)
                cell.text = str(row.get(h, ''))[:100]
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        self._set_arabic_font(run, 9)
        
        if len(data) > max_rows:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'... و {len(data) - max_rows} صف إضافي')
            self._set_arabic_font(run, 10)
            run.italic = True
        
        doc.add_paragraph()
    
    def _set_rtl(self, paragraph):
        """Set paragraph to RTL."""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    
    def _set_arabic_font(self, run, size=12):
        """Set Arabic font for a run."""
        if not run:
            return
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
