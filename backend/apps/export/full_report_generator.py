"""
Full Report Generator - Generate complete annual report as single document.
"""

import os
import io
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .chart_generator import generate_chart_image


class FullReportGenerator:
    """
    Generate complete annual report as single Word document.
    
    Usage:
        generator = FullReportGenerator(project)
        doc_path = generator.generate()
    """
    
    def __init__(self, project, output_dir: str = "./output"):
        self.project = project
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc = None
        self.stats = {
            'axes': 0,
            'items': 0,
            'texts': 0,
            'tables': 0,
            'figures': 0,
        }
    
    def generate(self, include_toc: bool = True, formats: str = 'all') -> dict:
        """
        Generate full report document.

        Uses project-specific data when available:
        - GeneratedContent for AI-generated paragraph text
        - TableData for contributor-submitted table rows
        - ItemStructure for modified component ordering
        Falls back to template defaults when project data is not available.

        Args:
            formats: 'html', 'docx', 'pdf', or 'all'

        Returns:
            Dict with paths: {'html': '...', 'docx': '...', 'pdf': '...'}
        """
        from apps.templates_app.models import Axis, Item, ItemComponent
        from apps.reports.models import ItemStructure, GeneratedContent, TableData

        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        results = {}

        # Pre-load project-specific data for efficient lookups
        self.generated_contents = {}  # {(item_id, component_id): GeneratedContent}
        self.table_data = {}  # {(item_id, table_def_id): TableData}

        if self.project:
            # Load all GeneratedContent for this project
            for gc in GeneratedContent.objects.filter(project=self.project).select_related('item_structure__item'):
                key = (gc.item_structure.item_id, gc.component_id)
                self.generated_contents[key] = gc

            # Load all TableData for this project
            for td in TableData.objects.filter(project=self.project).select_related('table_definition'):
                key = (td.table_definition.id if td.table_definition else None,)
                self.table_data[key] = td

        # Collect all content — FILTERED BY PROJECT TEMPLATE
        self.content = []
        if self.project and self.project.template:
            axes = Axis.objects.filter(template=self.project.template).order_by('order')
        else:
            axes = Axis.objects.all().order_by('order')

        for axis in axes:
            self.stats['axes'] += 1
            axis_content = {'type': 'axis', 'data': axis, 'items': []}

            items = sorted(Item.objects.filter(axis=axis), key=lambda i: [int(x) for x in i.code.split('.') if x.isdigit()])
            for item in items:
                self.stats['items'] += 1
                components = ItemComponent.objects.filter(item=item).order_by('order')
                axis_content['items'].append({
                    'item': item,
                    'components': list(components)
                })

            self.content.append(axis_content)
        
        # Generate HTML
        if formats in ('html', 'all'):
            html_path = self.output_dir / f"التقرير_السنوي_{timestamp}.html"
            self._generate_html(html_path)
            results['html'] = str(html_path)
        
        # Generate DOCX
        if formats in ('docx', 'all'):
            self.doc = Document()
            self._setup_document()
            self._add_title_page()
            if include_toc:
                self._add_toc_placeholder()

            for axis_content in self.content:
                self._add_axis(axis_content['data'])
                for item_data in axis_content['items']:
                    self._add_item_from_data(item_data)

            docx_path = self.output_dir / f"التقرير_السنوي_{timestamp}.docx"
            self.doc.save(str(docx_path))
            results['docx'] = str(docx_path)

        # Generate PDF (from HTML via WeasyPrint)
        if formats in ('pdf', 'all'):
            # Ensure HTML exists first
            html_path = self.output_dir / f"التقرير_السنوي_{timestamp}.html"
            if not html_path.exists():
                self._generate_html(html_path)
            try:
                import weasyprint
                pdf_path = self.output_dir / f"التقرير_السنوي_{timestamp}.pdf"
                weasyprint.HTML(filename=str(html_path)).write_pdf(str(pdf_path))
                results['pdf'] = str(pdf_path)
            except Exception as e:
                import logging
                logging.getLogger(__name__).error(f"PDF generation failed: {e}")

        return results
    
    def _generate_html(self, output_path: Path):
        """Generate full report as HTML."""
        project_name = self.project.name if self.project else 'التقرير السنوي'
        org_name = self.project.organization.name if self.project and self.project.organization else 'جامعة البترا'
        period = self.project.period if self.project else '2023-2024'

        html = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project_name}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Arial, sans-serif;
            direction: rtl;
            padding: 0;
            margin: 0;
            line-height: 1.9;
            color: #2d2d2d;
            background: #f4f4f4;
        }}
        .report {{ background: white; max-width: 1000px; margin: 0 auto; padding: 0; box-shadow: 0 2px 20px rgba(0,0,0,0.08); }}

        /* Title Page */
        .title-page {{
            text-align: center;
            padding: 80px 40px;
            background: linear-gradient(135deg, #f8f4f0 0%, #fff 50%, #f8f4f0 100%);
            border-bottom: 4px solid #8B1A1A;
            position: relative;
        }}
        .title-page::before {{
            content: '';
            position: absolute;
            top: 0; left: 0; right: 0;
            height: 6px;
            background: linear-gradient(to right, #8B1A1A, #C4A35A, #8B1A1A);
        }}
        .title-page .org-name {{ color: #8B1A1A; font-size: 1.8em; font-weight: 700; margin-bottom: 10px; }}
        .title-page h1 {{ color: #2d2d2d; font-size: 2.2em; margin: 20px 0; font-weight: 700; }}
        .title-page .period {{ color: #8B1A1A; font-size: 1.5em; font-weight: 600; }}
        .title-page .subtitle {{ color: #666; font-size: 0.95em; margin-top: 20px; line-height: 1.6; }}

        /* Content area */
        .content {{ padding: 40px; }}

        /* Axis (Section) */
        .axis {{ margin: 50px 0 30px; padding-top: 30px; border-top: 3px solid #8B1A1A; }}
        .axis-title {{ color: #8B1A1A; font-size: 1.4em; font-weight: 700; margin-bottom: 5px; }}
        .axis-subtitle {{ color: #C4A35A; font-size: 0.85em; font-weight: 600; margin-bottom: 20px; }}

        /* Item (Sub-section) */
        .item {{ margin: 25px 0; padding: 20px 25px; background: #fefefe; border-right: 4px solid #C4A35A; border-radius: 0 8px 8px 0; }}
        .item-title {{ color: #8B1A1A; font-size: 1.1em; font-weight: 700; margin-bottom: 12px; padding-bottom: 8px; border-bottom: 1px solid #eee; }}

        /* Text */
        .text {{ margin: 12px 0; text-align: justify; font-size: 0.95em; }}

        /* Tables */
        .table-container {{ margin: 20px 0; overflow-x: auto; }}
        .table-title {{ font-weight: 700; text-align: center; margin-bottom: 8px; color: #8B1A1A; font-size: 0.95em; }}
        table {{ width: 100%; border-collapse: collapse; font-size: 12.5px; }}
        th, td {{ border: 1px solid #d4d4d4; padding: 8px 10px; text-align: center; }}
        th {{ background: #8B1A1A; color: white; font-weight: 600; font-size: 12px; }}
        tr:nth-child(even) {{ background: #faf8f5; }}
        tr:hover {{ background: #f0ebe3; }}

        /* Figures */
        .figure {{ text-align: center; margin: 25px 0; }}
        .figure img {{ max-width: 100%; border-radius: 4px; border: 1px solid #eee; }}
        .figure-title {{ font-weight: 600; color: #8B1A1A; margin-top: 8px; font-size: 0.9em; }}

        /* TOC */
        .toc {{ margin: 0; padding: 30px 40px; background: #faf8f5; border-bottom: 2px solid #e8e0d4; }}
        .toc h3 {{ color: #8B1A1A; font-size: 1.2em; margin-bottom: 15px; }}
        .toc ul {{ list-style: none; padding: 0; }}
        .toc li {{ padding: 4px 0; font-size: 0.9em; }}
        .toc li a {{ color: #444; text-decoration: none; }}
        .toc li a:hover {{ color: #8B1A1A; }}
        .toc li strong {{ color: #8B1A1A; }}

        /* Footer */
        .footer {{ text-align: center; margin-top: 50px; padding: 25px 40px; color: #999; font-size: 11px; border-top: 2px solid #e8e0d4; background: #faf8f5; }}

        @media print {{
            body {{ background: white; padding: 0; }}
            .report {{ box-shadow: none; }}
            .title-page {{ page-break-after: always; }}
            .axis {{ page-break-before: always; }}
        }}
    </style>
</head>
<body>
    <div class="report">
        <div class="title-page">
            <div class="org-name">{org_name}</div>
            <h1>التقرير السنوي عن أداء الجامعة</h1>
            <div class="period">للعام الجامعي ({period})</div>
            <div class="subtitle">حسب النموذج المعتمد في قرار مجلس التعليم العالي</div>
        </div>
        
        <div class="toc">
            <h3>فهرس المحتويات</h3>
            <ul>
'''
        # TOC
        for axis_content in self.content:
            axis = axis_content['data']
            html += f'<li><a href="#axis-{axis.order}"><strong>المحور {axis.order}: {axis.name}</strong></a></li>\n'
            for item_data in axis_content['items']:
                item = item_data['item']
                html += f'<li style="padding-right: 20px;"><a href="#item-{item.code.replace(".", "-")}">{item.code}: {item.name}</a></li>\n'
        
        html += '''
            </ul>
        </div>
        <div class="content">
'''

        # Content
        for axis_content in self.content:
            axis = axis_content['data']
            html += f'''
        <div class="axis" id="axis-{axis.order}">
            <h2 class="axis-title">المحور {axis.order}: {axis.name}</h2>
'''
            for item_data in axis_content['items']:
                item = item_data['item']
                html += f'''
            <div class="item" id="item-{item.code.replace(".", "-")}">
                <h3 class="item-title">{item.code}: {item.name}</h3>
'''
                for comp in item_data['components']:
                    html += self._component_to_html(comp, item=item)
                
                html += '            </div>\n'
            
            html += '        </div>\n'
        
        html += f'''
        </div><!-- end content -->
        <div class="footer">
            تم التوليد بواسطة نظام تقرير.ai<br>
            المحاور: {self.stats['axes']} | البنود: {self.stats['items']} | النصوص: {self.stats['texts']} | الجداول: {self.stats['tables']} | الأشكال: {self.stats['figures']}
        </div>
    </div>
</body>
</html>'''
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
    
    def _component_to_html(self, comp, item=None) -> str:
        """Convert component to HTML.

        Prefers project data (GeneratedContent, TableData) over template defaults.
        """
        config = comp.config or {}

        if comp.component_type in ('text', 'text_ai'):
            # 1. Try GeneratedContent (AI-generated or user-edited) first
            content = ''
            if item and comp.ref_id:
                gc = self.generated_contents.get((item.id, comp.ref_id))
                if gc and gc.status in ('generated', 'edited', 'approved'):
                    content = gc.manual_edit or gc.content

            # 2. Fall back to template's original text
            if not content:
                content = config.get('full_text', '') or config.get('preview', '')

            if content:
                self.stats['texts'] += 1
                return f'<p class="text">{content}</p>\n'

        elif comp.component_type == 'table':
            # 1. Try TableData (contributor-submitted) first
            headers = []
            data_rows = []
            if comp.table_ref_id:
                td = self.table_data.get((comp.table_ref_id,))
                if td and td.rows:
                    # TableData.rows format: list of dicts or lists
                    if td.table_definition and td.table_definition.columns:
                        headers = [c.get('name', c) if isinstance(c, dict) else c for c in td.table_definition.columns]
                    data_rows = td.rows

            # 2. Fall back to template's extracted_data
            if not (headers and data_rows):
                extracted = config.get('extracted_data', {})
                headers = extracted.get('headers', [])
                data_rows = extracted.get('data', [])

            if headers and data_rows:
                self.stats['tables'] += 1
                html = '<div class="table-container">\n'
                if comp.title:
                    html += f'<p class="table-title">{comp.title}</p>\n'
                html += '<table><thead><tr>'
                for h in headers:
                    html += f'<th>{h}</th>'
                html += '</tr></thead><tbody>\n'

                for row in data_rows[:100]:
                    html += '<tr>'
                    for j, h in enumerate(headers):
                        val = row[j] if isinstance(row, list) and j < len(row) else row.get(h, '') if isinstance(row, dict) else ''
                        html += f'<td>{val}</td>'
                    html += '</tr>\n'

                html += '</tbody></table>\n'
                if len(data_rows) > 100:
                    html += f'<p style="text-align:center;font-size:12px;color:#666;">... و {len(data_rows)-100} صف إضافي</p>\n'
                html += '</div>\n'
                return html
        
        elif comp.component_type in ('figure', 'chart'):
            self.stats['figures'] += 1
            title = comp.title or 'شكل'
            
            # Check for static image first
            static_image = config.get('static_image')
            if static_image:
                return f'<div class="figure"><img src="{static_image}" alt="{title}"><p class="figure-title">{title}</p></div>\n'
            
            # Try to generate chart from data
            chart_data = config.get('chart_data', {})
            if chart_data:
                series = chart_data.get('series', [])
                if series and isinstance(series[0], dict):
                    labels = series[0].get('categories', []) or chart_data.get('categories', [])
                    values = series[0].get('values', [])
                    
                    if labels and values:
                        # Generate chart image
                        chart_type = (chart_data.get('type') or 'bar').replace('Chart', '')
                        img_config = {
                            'type': chart_type,
                            'title': title,
                            'data': {
                                'labels': labels,
                                'datasets': [{'label': '', 'values': values}]
                            }
                        }
                        try:
                            img_buffer = generate_chart_image(img_config)
                            if img_buffer:
                                img_name = f"chart_{self.stats['figures']}.png"
                                img_path = self.output_dir / img_name
                                with open(img_path, 'wb') as f:
                                    f.write(img_buffer.read())
                                return f'<div class="figure"><img src="{img_name}" alt="{title}"><p class="figure-title">{title}</p></div>\n'
                        except Exception as e:
                            pass
            
            return f'<div class="figure"><p class="figure-title">[{title}]</p></div>\n'
        
        return ''
    
    def _add_item_from_data(self, item_data):
        """Add item with all its components from collected data."""
        item = item_data['item']
        
        # Item heading
        heading = self.doc.add_heading(f"{item.code}: {item.name}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_rtl(heading)
        
        for run in heading.runs:
            self._set_arabic_font(run, 14)
        
        for comp in item_data['components']:
            if comp.component_type == 'text':
                self._add_text_component(comp)
            elif comp.component_type == 'table':
                self._add_table_component(comp)
            elif comp.component_type in ('figure', 'chart'):
                self._add_figure_component(comp)
        
        self.doc.add_paragraph()
    
    def _setup_document(self):
        """Setup document margins and styles."""
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
    
    def _add_title_page(self):
        """Add title page."""
        # University name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("جامعة البترا")
        self._set_arabic_font(run, 28)
        run.bold = True
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Report title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("التقرير السنوي")
        self._set_arabic_font(run, 36)
        run.bold = True
        run.font.color.rgb = RGBColor(26, 95, 122)
        
        self.doc.add_paragraph()
        
        # Year
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("2023-2024")
        self._set_arabic_font(run, 32)
        run.bold = True
        
        # Page break
        self.doc.add_page_break()
    
    def _add_toc_placeholder(self):
        """Add table of contents placeholder."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("فهرس المحتويات")
        self._set_arabic_font(run, 18)
        run.bold = True
        
        self.doc.add_paragraph()
        
        # Note about TOC
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[يتم تحديث الفهرس تلقائياً في Word: References → Update Table]")
        self._set_arabic_font(run, 11)
        run.italic = True
        
        self.doc.add_page_break()
    
    def _add_axis(self, axis):
        """Add axis heading."""
        # Page break before new axis (except first)
        if self.stats['axes'] > 1:
            self.doc.add_page_break()
        
        # Axis title
        heading = self.doc.add_heading(f"المحور {axis.order}: {axis.name}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_rtl(heading)
        
        for run in heading.runs:
            self._set_arabic_font(run, 18)
            run.font.color.rgb = RGBColor(26, 95, 122)
        
        self.doc.add_paragraph()
    
    def _add_item(self, item):
        """Add item with all its components."""
        from apps.templates_app.models import ItemComponent
        
        # Item heading
        heading = self.doc.add_heading(f"{item.code}: {item.name}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_rtl(heading)
        
        for run in heading.runs:
            self._set_arabic_font(run, 14)
        
        # Get components ordered
        components = ItemComponent.objects.filter(item=item).order_by('order')
        
        for comp in components:
            if comp.component_type == 'text':
                self._add_text_component(comp)
            elif comp.component_type == 'table':
                self._add_table_component(comp)
            elif comp.component_type in ('figure', 'chart'):
                self._add_figure_component(comp)
        
        self.doc.add_paragraph()
    
    def _add_text_component(self, comp):
        """Add text paragraph."""
        config = comp.config or {}
        content = config.get('full_text', '') or config.get('preview', '')
        
        if not content:
            return
        
        self.stats['texts'] += 1
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_rtl(p)
        
        run = p.add_run(content)
        self._set_arabic_font(run, 12)
    
    def _add_table_component(self, comp):
        """Add table with data."""
        config = comp.config or {}
        extracted = config.get('extracted_data', {})
        
        if not extracted:
            return
        
        headers = extracted.get('headers', [])
        data_rows = extracted.get('data', [])
        
        if not headers or not data_rows:
            return
        
        self.stats['tables'] += 1
        
        # Table title
        if comp.title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(comp.title)
            self._set_arabic_font(run, 11)
            run.bold = True
        
        # Create table
        num_rows = min(len(data_rows), 100) + 1  # Limit rows + header
        num_cols = len(headers)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        header_row = table.rows[0]
        for j, h in enumerate(headers):
            cell = header_row.cells[j]
            cell.text = str(h)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    self._set_arabic_font(run, 9)
                    run.bold = True
        
        # Data rows
        for i, row in enumerate(data_rows[:100]):
            if i + 1 >= num_rows:
                break
            table_row = table.rows[i + 1]
            for j, h in enumerate(headers):
                if j >= num_cols:
                    break
                cell = table_row.cells[j]
                val = row[j] if isinstance(row, list) and j < len(row) else row.get(h, '') if isinstance(row, dict) else ''
                cell.text = str(val)[:100]
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        self._set_arabic_font(run, 9)
        
        # Note if truncated
        if len(data_rows) > 100:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"... و {len(data_rows) - 100} صف إضافي")
            self._set_arabic_font(run, 9)
            run.italic = True
        
        self.doc.add_paragraph()
    
    def _add_figure_component(self, comp):
        """Add figure/chart."""
        config = comp.config or {}
        
        self.stats['figures'] += 1
        
        # Check for static image first
        static_image = config.get('static_image')
        if static_image:
            img_path = self.output_dir / static_image
            if img_path.exists():
                self.doc.add_picture(str(img_path), width=Inches(5))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if comp.title:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(comp.title)
                    self._set_arabic_font(run, 10)
                    run.italic = True
                
                self.doc.add_paragraph()
                return
        
        chart_data = config.get('chart_data', {})
        
        if not chart_data:
            return
        
        # Try to generate chart image
        try:
            series = chart_data.get('series', [])
            
            # Labels might be inside series[0] OR at top level
            if series and isinstance(series[0], dict):
                labels = series[0].get('categories', []) or chart_data.get('categories', []) or chart_data.get('labels', [])
                values = series[0].get('values', series[0].get('data', []))
            elif series:
                labels = chart_data.get('categories', []) or chart_data.get('labels', [])
                values = series
            else:
                labels = chart_data.get('categories', []) or chart_data.get('labels', [])
                values = []
            
            if labels and values:
                chart_type = (chart_data.get('type') or 'bar').replace('Chart', '')
                img_config = {
                    'type': chart_type,
                    'title': comp.title or '',
                    'data': {
                        'labels': labels,
                        'datasets': [{'label': '', 'values': values}]
                    }
                }
                
                img_buffer = generate_chart_image(img_config)
                
                if img_buffer:
                    # Save temp image
                    img_path = self.output_dir / f"temp_chart_{self.stats['figures']}.png"
                    with open(img_path, 'wb') as f:
                        f.write(img_buffer.read())
                    
                    # Add to document
                    self.doc.add_picture(str(img_path), width=Inches(5))
                    self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Clean up
                    os.remove(img_path)
        except Exception as e:
            print(f"Warning: Could not generate chart: {e}")
        
        # Figure caption
        if comp.title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(comp.title)
            self._set_arabic_font(run, 10)
            run.italic = True
        
        self.doc.add_paragraph()
    
    def _set_rtl(self, paragraph):
        """Set paragraph to RTL."""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    
    def _set_arabic_font(self, run, size=12):
        """Set Arabic font."""
        if not run:
            return
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    
    def get_stats(self) -> dict:
        """Get generation statistics."""
        return self.stats.copy()
