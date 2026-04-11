"""
Report Generator - Fixed version that reads from ItemComponent model.

Flow:
1. Load ItemComponents from DB (ordered)
2. Build HTML with actual data
3. Export to DOCX/PDF
"""

import os
import io
import json
from pathlib import Path
from typing import Optional, Dict, List, Any
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .chart_generator import generate_chart_image


@dataclass
class Component:
    """A single component in the report."""
    id: str
    type: str  # text, table, figure
    title: str = ""
    content: str = ""
    data: Any = None
    image_path: str = ""
    status: str = "pending"


@dataclass 
class GeneratedReport:
    """Generated report output."""
    item_code: str
    item_title: str
    components: List[Component] = field(default_factory=list)
    html_path: str = ""
    docx_path: str = ""
    pdf_path: str = ""


class ReportGenerator:
    """
    Report generation using ItemComponent model.
    
    Usage:
        generator = ReportGenerator(project, output_dir)
        result = generator.generate_item(item_structure, formats='all')
    """
    
    def __init__(self, project, output_dir: str = "./output"):
        self.project = project
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_item(self, item_structure, formats: str = 'all') -> Dict[str, str]:
        """
        Generate report for an item.
        
        Args:
            item_structure: ItemStructure model instance
            formats: 'html', 'docx', 'pdf', or 'all'
        
        Returns:
            Dict with paths: {'html': '...', 'docx': '...', 'pdf': '...'}
        """
        item = item_structure.item
        item_code = item.code.replace('.', '_')
        
        # Build components from ItemComponent model (the actual data!)
        components = self._build_components_from_db(item)
        
        # Create report object
        report = GeneratedReport(
            item_code=item.code,
            item_title=item.name,
            components=components
        )
        
        results = {}
        
        # Generate requested formats
        if formats in ('html', 'all'):
            html_path = self.output_dir / f"item_{item_code}.html"
            self._generate_html(report, html_path)
            results['html'] = str(html_path)
        
        if formats in ('docx', 'all'):
            docx_path = self.output_dir / f"item_{item_code}.docx"
            self._generate_docx(report, docx_path)
            results['docx'] = str(docx_path)
        
        if formats in ('pdf', 'all'):
            pdf_path = self.output_dir / f"item_{item_code}.pdf"
            if self._generate_pdf(report, pdf_path):
                results['pdf'] = str(pdf_path)
            else:
                results['pdf'] = None
        
        return results
    
    def _build_components_from_db(self, item) -> List[Component]:
        """
        Build components from ItemComponent model - THE ACTUAL DATA SOURCE.
        """
        from apps.templates_app.models import ItemComponent as ItemComponentModel
        
        components = []
        
        # Get all ItemComponents ordered by 'order' field
        db_components = ItemComponentModel.objects.filter(item=item).order_by('order')
        
        for idx, db_comp in enumerate(db_components):
            comp_id = f"{db_comp.component_type[0]}{idx + 1}"  # t1, p1, f1, etc.
            
            comp = Component(
                id=comp_id,
                type=db_comp.component_type,
                title=db_comp.title or "",
            )
            
            config = db_comp.config or {}
            
            if db_comp.component_type == 'text':
                # Get full text from config
                comp.content = config.get('full_text', '') or config.get('preview', '')
                comp.status = 'loaded' if comp.content else 'empty'
            
            elif db_comp.component_type in ('figure', 'chart'):
                comp.type = 'figure'  # Normalize type
                
                # Check for static image first (extracted from original)
                static_image = config.get('static_image')
                if static_image:
                    img_path = self.output_dir / static_image
                    if img_path.exists():
                        comp.image_path = str(img_path)
                        comp.status = 'static'
                    else:
                        comp.status = 'no_image'
                else:
                    # Try to generate from chart data
                    chart_data = config.get('chart_data', {})
                    if chart_data:
                        comp.data = chart_data
                        img_path = self._generate_chart_for_component(comp, chart_data)
                        if img_path:
                            comp.image_path = img_path
                            comp.status = 'generated'
                        else:
                            comp.status = 'no_image'
                    else:
                        comp.status = 'no_data'
            
            elif db_comp.component_type == 'table':
                # Get extracted data from config
                extracted = config.get('extracted_data', {})
                if extracted:
                    # Convert to list format expected by HTML generator
                    headers = extracted.get('headers', [])
                    data_rows = extracted.get('data', [])
                    
                    if headers and data_rows:
                        # Convert to list of dicts
                        comp.data = []
                        for row in data_rows:
                            if isinstance(row, list):
                                row_dict = {}
                                for i, header in enumerate(headers):
                                    row_dict[header] = row[i] if i < len(row) else ''
                                comp.data.append(row_dict)
                            elif isinstance(row, dict):
                                comp.data.append(row)
                        
                        comp.status = 'loaded'
                    else:
                        comp.status = 'no_data'
                else:
                    comp.status = 'no_data'
            
            components.append(comp)
        
        return components
    
    def _generate_chart_for_component(self, comp: Component, chart_data: dict) -> Optional[str]:
        """Generate chart image from chart_data."""
        try:
            # Extract labels and values from chart_data
            labels = []
            values = []
            chart_type = chart_data.get('type', 'bar').replace('Chart', '')  # barChart -> bar
            
            # Handle different chart_data formats
            series = chart_data.get('series', [])
            
            if series and isinstance(series, list) and isinstance(series[0], dict):
                # Categories might be inside series[0] OR at top level
                labels = series[0].get('categories', []) or chart_data.get('categories', []) or chart_data.get('labels', [])
                values = series[0].get('values', series[0].get('data', []))
            elif series and isinstance(series, list):
                values = series
                labels = chart_data.get('categories', []) or chart_data.get('labels', [])
            else:
                labels = chart_data.get('categories', []) or chart_data.get('labels', [])
                values = chart_data.get('values', [])
            
            if not labels or not values:
                return None
            
            # Generate chart
            img_config = {
                'type': chart_type,
                'title': comp.title,
                'data': {
                    'labels': labels,
                    'datasets': [{'label': '', 'values': values}]
                }
            }
            
            img_buffer = generate_chart_image(img_config)
            
            if img_buffer:
                img_path = self.output_dir / f"{comp.id}.png"
                with open(img_path, 'wb') as f:
                    f.write(img_buffer.read())
                return str(img_path)
        
        except Exception as e:
            print(f"Warning: Could not generate chart for {comp.id}: {e}")
        
        return None
    
    def _generate_html(self, report: GeneratedReport, output_path: Path):
        """Generate HTML file."""
        html = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report.item_code}: {report.item_title}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{ 
            font-family: 'Segoe UI', Tahoma, Arial, sans-serif;
            direction: rtl;
            padding: 40px;
            max-width: 900px;
            margin: 0 auto;
            line-height: 1.8;
            color: #333;
        }}
        h1 {{ 
            color: #1a5f7a;
            border-bottom: 3px solid #1a5f7a;
            padding-bottom: 10px;
        }}
        .component {{ margin: 25px 0; }}
        .text p {{ 
            text-align: justify;
            margin: 15px 0;
        }}
        .figure {{ 
            text-align: center;
            margin: 30px 0;
        }}
        .figure img {{ 
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        .figure-title {{
            font-style: italic;
            color: #666;
            margin-top: 10px;
        }}
        .table {{ 
            margin: 30px 0;
            overflow-x: auto;
        }}
        .table-title {{
            font-weight: bold;
            text-align: center;
            margin-bottom: 10px;
        }}
        table {{ 
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
        }}
        th, td {{ 
            border: 1px solid #ddd;
            padding: 10px;
            text-align: center;
        }}
        th {{ 
            background: #1a5f7a;
            color: white;
            font-weight: bold;
        }}
        tr:nth-child(even) {{ background: #f9f9f9; }}
        tr:hover {{ background: #f0f7fa; }}
        .placeholder {{
            background: #fff3cd;
            border: 1px dashed #ffc107;
            padding: 15px;
            border-radius: 5px;
            color: #856404;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
            font-size: 12px;
            text-align: center;
        }}
    </style>
</head>
<body>
    <article>
        <h1>{report.item_code}: {report.item_title}</h1>
'''
        
        for comp in report.components:
            html += self._component_to_html(comp)
        
        html += '''
        <div class="footer">
            تم التوليد بواسطة نظام تقرير.ai
        </div>
    </article>
</body>
</html>'''
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
    
    def _component_to_html(self, comp: Component) -> str:
        """Convert component to HTML."""
        if comp.type == 'text':
            if comp.content:
                return f'''
        <div class="component text" data-id="{comp.id}">
            <p>{comp.content}</p>
        </div>'''
            else:
                return f'''
        <div class="component text placeholder" data-id="{comp.id}">
            <p>[{comp.title or comp.id} - في انتظار النص]</p>
        </div>'''
        
        elif comp.type == 'figure':
            if comp.image_path and os.path.exists(comp.image_path):
                img_name = os.path.basename(comp.image_path)
                return f'''
        <div class="component figure" data-id="{comp.id}">
            <img src="{img_name}" alt="{comp.title}">
            <p class="figure-title">{comp.title}</p>
        </div>'''
            else:
                return f'''
        <div class="component figure placeholder" data-id="{comp.id}">
            <p>[{comp.title} - في انتظار الرسم البياني]</p>
        </div>'''
        
        elif comp.type == 'table':
            if comp.data:
                return self._table_to_html(comp)
            else:
                return f'''
        <div class="component table placeholder" data-id="{comp.id}">
            <p>[{comp.title} - في انتظار البيانات]</p>
        </div>'''
        
        return ''
    
    def _table_to_html(self, comp: Component) -> str:
        """Convert table data to HTML."""
        data = comp.data
        if not data or not isinstance(data, list):
            return ''
        
        html = f'''
        <div class="component table" data-id="{comp.id}">
            <p class="table-title">{comp.title}</p>
            <table>'''
        
        if data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            
            # Header row
            html += '\n                <thead><tr>'
            for h in headers:
                html += f'<th>{h}</th>'
            html += '</tr></thead>'
            
            # Data rows (limit to 100)
            html += '\n                <tbody>'
            for row in data[:100]:
                html += '\n                    <tr>'
                for h in headers:
                    val = row.get(h, '')
                    html += f'<td>{val}</td>'
                html += '</tr>'
            html += '\n                </tbody>'
            
            if len(data) > 100:
                html += f'\n                <tfoot><tr><td colspan="{len(headers)}">... و {len(data) - 100} صف إضافي</td></tr></tfoot>'
        
        html += '''
            </table>
        </div>'''
        
        return html
    
    def _generate_docx(self, report: GeneratedReport, output_path: Path):
        """Generate Word document."""
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Title
        title = doc.add_heading(f"{report.item_code}: {report.item_title}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_arabic_font(title.runs[0] if title.runs else None, 18)
        
        doc.add_paragraph()
        
        # Components
        for comp in report.components:
            self._add_component_to_docx(doc, comp)
        
        doc.save(str(output_path))
    
    def _add_component_to_docx(self, doc: Document, comp: Component):
        """Add component to Word document."""
        if comp.type == 'text' and comp.content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._set_rtl(p)
            run = p.add_run(comp.content)
            self._set_arabic_font(run, 12)
        
        elif comp.type == 'figure' and comp.image_path and os.path.exists(comp.image_path):
            doc.add_picture(comp.image_path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if comp.title:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(comp.title)
                self._set_arabic_font(run, 10)
                run.italic = True
            
            doc.add_paragraph()
        
        elif comp.type == 'table' and comp.data:
            self._add_table_to_docx(doc, comp)
    
    def _add_table_to_docx(self, doc: Document, comp: Component):
        """Add table to Word document."""
        data = comp.data
        if not data or not isinstance(data, list) or not isinstance(data[0], dict):
            return
        
        # Title
        if comp.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(comp.title)
            self._set_arabic_font(run, 12)
            run.bold = True
        
        headers = list(data[0].keys())
        max_rows = min(len(data), 50)
        
        table = doc.add_table(rows=max_rows + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Headers
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = str(h)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    self._set_arabic_font(run, 9)
                    run.bold = True
        
        # Data
        for i, row in enumerate(data[:max_rows]):
            for j, h in enumerate(headers):
                cell = table.cell(i + 1, j)
                cell.text = str(row.get(h, ''))[:100]
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        self._set_arabic_font(run, 9)
        
        if len(data) > max_rows:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'... و {len(data) - max_rows} صف إضافي')
            self._set_arabic_font(run, 10)
            run.italic = True
        
        doc.add_paragraph()
    
    def _generate_pdf(self, report: GeneratedReport, output_path: Path) -> bool:
        """Generate PDF from HTML."""
        try:
            import weasyprint
            
            html_path = output_path.with_suffix('.html')
            self._generate_html(report, html_path)
            
            weasyprint.HTML(filename=str(html_path)).write_pdf(str(output_path))
            return True
        except ImportError:
            try:
                import subprocess
                html_path = output_path.with_suffix('.html')
                self._generate_html(report, html_path)
                
                result = subprocess.run(
                    ['wkhtmltopdf', str(html_path), str(output_path)],
                    capture_output=True
                )
                return result.returncode == 0
            except:
                print("Warning: Could not generate PDF.")
                return False
        except Exception as e:
            print(f"Warning: PDF generation failed: {e}")
            return False
    
    def _set_rtl(self, paragraph):
        """Set paragraph to RTL."""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)
    
    def _set_arabic_font(self, run, size=12):
        """Set Arabic font."""
        if not run:
            return
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
