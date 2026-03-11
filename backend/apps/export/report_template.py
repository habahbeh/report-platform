"""
Professional Report Template Generator
يولد تقارير Word بنفس شكل التقرير السنوي الأصلي لجامعة البترا
"""

import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class PetraReportGenerator:
    """Generate reports matching Petra University annual report style."""
    
    def __init__(self, period):
        self.period = period
        self.doc = Document()
        self.table_counter = 0
        self.figure_counter = 0
        self.current_axis_code = "1"
        self.setup_document()
    
    def setup_document(self):
        """Configure document like original report."""
        # Page margins (matching original)
        for section in self.doc.sections:
            section.top_margin = Cm(1.7)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(1.6)
            section.right_margin = Cm(2.5)
        
        self._create_styles()
    
    def _create_styles(self):
        """Create styles matching original report."""
        styles = self.doc.styles
        
        # Main Title - 30pt
        if 'PetraTitle' not in [s.name for s in styles]:
            style = styles.add_style('PetraTitle', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(30)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.space_after = Pt(6)
        
        # Heading 1 - المعيار (18pt)
        if 'PetraH1' not in [s.name for s in styles]:
            style = styles.add_style('PetraH1', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(18)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(12)
        
        # Heading 2 - المؤشر (16pt)
        if 'PetraH2' not in [s.name for s in styles]:
            style = styles.add_style('PetraH2', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(16)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 128)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        
        # Heading 3 - عنوان جدول/شكل (14pt)
        if 'PetraH3' not in [s.name for s in styles]:
            style = styles.add_style('PetraH3', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(14)
            style.font.bold = True
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        
        # Normal body text (12pt)
        if 'PetraBody' not in [s.name for s in styles]:
            style = styles.add_style('PetraBody', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.first_line_indent = Cm(0.5)
    
    def _set_rtl(self, paragraph):
        """Set paragraph to RTL."""
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    
    def _add_paragraph(self, text, style='PetraBody'):
        """Add RTL paragraph with style."""
        p = self.doc.add_paragraph(text, style=style)
        self._set_rtl(p)
        return p
    
    def add_cover_page(self):
        """Add cover page like original."""
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Main title
        p = self._add_paragraph("التـّقـريـر السّـنـويّ عن أداء الجامعة", 'PetraTitle')
        
        # Subtitle
        p = self._add_paragraph(f"للعام الجامعيّ ({self.period.academic_year})", 'PetraTitle')
        p.runs[0].font.size = Pt(28)
        
        self.doc.add_paragraph()
        
        # Reference text
        p = self._add_paragraph("حسب النموذج المعتمد في قرار مجلس التعليم العالي")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(16)
        
        for _ in range(4):
            self.doc.add_paragraph()
        
        # From/To
        p = self._add_paragraph("مُقَدَّم من رئيس الجامعة", 'PetraTitle')
        p.runs[0].font.size = Pt(24)
        
        p = self._add_paragraph("إلى مَجْلِس الأمناء", 'PetraTitle')
        p.runs[0].font.size = Pt(24)
        
        for _ in range(6):
            self.doc.add_paragraph()
        
        # Date
        p = self._add_paragraph(f"تاريخ الإصدار: {datetime.now().strftime('%Y/%m/%d')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents."""
        p = self._add_paragraph("فهرس المحتويات", 'PetraH1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # List axes
        axes = self.period.template.axes.all().order_by('order')
        for axis in axes:
            p = self._add_paragraph(f"المعيار {axis.code}: {axis.name}")
            p.runs[0].font.size = Pt(14)
        
        self.doc.add_page_break()
    
    def add_axis(self, axis_draft):
        """Add axis section (المعيار)."""
        axis = axis_draft.axis
        self.current_axis_code = axis.code
        self.table_counter = 0
        self.figure_counter = 0
        
        # Axis heading
        p = self._add_paragraph(f"المعيار {axis.code}: {axis.name}", 'PetraH1')
        
        # Axis general content (if any)
        if axis_draft.content:
            content = self._clean_content(axis_draft.content)
            self._add_body_text(content)
    
    def add_item(self, item_draft):
        """Add item section (المؤشر) with text, table, and chart."""
        item = item_draft.item
        
        # Item heading (like original: "1.1: عدد البرامج...")
        p = self._add_paragraph(f"{item.code}: {item.name}:", 'PetraH2')
        
        # Body text
        if item_draft.content:
            content = self._clean_content(item_draft.content)
            self._add_body_text(content)
        
        # Table
        if item_draft.table_data:
            self.table_counter += 1
            table_title = f"جدول ({self.current_axis_code}-{self.table_counter}): {item.name}"
            self._add_table(item_draft.table_data, table_title)
        
        # Chart
        if item_draft.chart_config:
            self.figure_counter += 1
            figure_title = f"شكل ({self.current_axis_code}-{self.figure_counter}): {item.name}"
            self._add_chart(item_draft.chart_config, figure_title)
    
    def _clean_content(self, content: str) -> str:
        """Remove tags from content."""
        content = re.sub(r'\[TABLE\].*?\[/TABLE\]', '', content, flags=re.DOTALL)
        content = re.sub(r'\[CHART\].*?\[/CHART\]', '', content, flags=re.DOTALL)
        content = re.sub(r'\[IMAGE\].*?\[/IMAGE\]', '', content, flags=re.DOTALL)
        content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
        content = re.sub(r'\n---+\n', '\n\n', content)
        content = re.sub(r'\n{3,}', '\n\n', content)
        return content.strip()
    
    def _add_body_text(self, content: str):
        """Add body text as paragraphs."""
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text and len(para_text) > 10:
                p = self._add_paragraph(para_text, 'PetraBody')
    
    def _add_table(self, table_data: list, title: str):
        """Add formatted table like original report."""
        if not table_data:
            return
        
        # Filter separator rows
        clean_data = []
        for row in table_data:
            if isinstance(row, dict):
                values = list(row.values())
                if any(v and ':' not in str(v)[:3] and '-' not in str(v)[:3] for v in values):
                    clean_data.append(row)
        
        if not clean_data:
            return
        
        # Table title
        p = self._add_paragraph(title, 'PetraH3')
        
        headers = list(clean_data[0].keys())
        
        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # RTL table
        tblPr = table._tbl.tblPr
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)
        
        # Header row styling
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_rtl(cell.paragraphs[0])
            
            # Bold white text on dark blue background
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Dark blue background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '003366')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # Data rows
        for row_data in clean_data:
            row = table.add_row()
            for i, header in enumerate(headers):
                cell = row.cells[i]
                value = row_data.get(header, '')
                cell.text = str(value) if value else ''
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_rtl(cell.paragraphs[0])
                for run in cell.paragraphs[0].runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        
        self.doc.add_paragraph()
    
    def _add_chart(self, chart_config: dict, title: str):
        """Add chart image with title like original."""
        if not chart_config:
            return
        
        try:
            from .chart_generator import generate_chart_image
            
            chart_buffer = generate_chart_image(chart_config)
            
            if chart_buffer:
                # Add figure title BEFORE image (like original)
                p = self._add_paragraph(title, 'PetraH3')
                
                # Add image
                self.doc.add_picture(chart_buffer, width=Inches(5))
                last_paragraph = self.doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                self.doc.add_paragraph()
                
        except Exception as e:
            print(f"Chart error: {e}")
    
    def add_footer(self):
        """Add footer."""
        self.doc.add_page_break()
        
        p = self._add_paragraph("ـــــــــــــــــــــــــــــــــــــــــــ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        
        p = self._add_paragraph("تم توليد هذا التقرير آلياً بواسطة منصة تقرير.ai")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        
        p = self._add_paragraph(f"بتاريخ: {datetime.now().strftime('%Y-%m-%d')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    def save(self) -> io.BytesIO:
        """Save document."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer


def generate_professional_report(period_or_project, generated_report=None, include_items=True) -> io.BytesIO:
    """
    Generate professional report matching Petra University style.
    """
    from apps.data_collection.models import DataCollectionPeriod
    from apps.reports.models import AxisDraft, ItemDraft
    
    def update_progress(percent, step):
        if generated_report:
            generated_report.progress = percent
            generated_report.current_step = step
            generated_report.save(update_fields=['progress', 'current_step'])
    
    # Get period
    if hasattr(period_or_project, 'template') and hasattr(period_or_project, 'period'):
        period = DataCollectionPeriod.objects.filter(template=period_or_project.template).first()
    elif isinstance(period_or_project, DataCollectionPeriod):
        period = period_or_project
    else:
        period = DataCollectionPeriod.objects.first()
    
    if not period:
        raise ValueError("لا توجد فترة جمع")
    
    update_progress(5, 'تهيئة المستند')
    
    generator = PetraReportGenerator(period)
    
    update_progress(10, 'إنشاء صفحة الغلاف')
    generator.add_cover_page()
    
    update_progress(15, 'إنشاء الفهرس')
    generator.add_table_of_contents()
    
    # Get axes
    axes = period.template.axes.all().order_by('order')
    total = axes.count()
    
    for i, axis in enumerate(axes):
        progress = 20 + int((i / total) * 70)
        update_progress(progress, f'المعيار {axis.code}: {axis.name}')
        
        # Get or create axis draft
        axis_draft, _ = AxisDraft.objects.get_or_create(
            period=period, axis=axis,
            defaults={'status': 'not_started', 'content': ''}
        )
        
        generator.add_axis(axis_draft)
        
        # Add items
        if include_items:
            item_drafts = ItemDraft.objects.filter(
                period=period, item__axis=axis
            ).select_related('item').order_by('item__order')
            
            for item_draft in item_drafts:
                if item_draft.content or item_draft.table_data or item_draft.chart_config:
                    generator.add_item(item_draft)
    
    update_progress(92, 'إضافة التذييل')
    generator.add_footer()
    
    update_progress(95, 'حفظ الملف')
    buffer = generator.save()
    
    update_progress(100, 'اكتمل')
    
    return buffer
