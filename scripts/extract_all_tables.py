#!/usr/bin/env python3
"""
Extract all tables from the original Word report.
"""

import json
import os
import re
import sys
from pathlib import Path

# Add backend to path for docx
sys.path.insert(0, str(Path(__file__).parent.parent / 'backend'))

from docx import Document

# Paths
REPORT_PATH = Path(__file__).parent.parent / "habahbeh/التقرير السنوي لجامعة البترا 2023-2024 حتى تاريخ 29.09.2025.docx"
DATA_DIR = Path(__file__).parent.parent / "data"
OUTPUT_FILE = DATA_DIR / "extracted_tables.json"


def extract_table_data(table):
    """Extract data from a Word table."""
    data = []
    headers = []
    
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            text = cell.text.strip()
            row_data.append(text)
        
        if i == 0:
            headers = row_data
        else:
            if headers:
                row_dict = {}
                for j, val in enumerate(row_data):
                    if j < len(headers):
                        row_dict[headers[j]] = val
                    else:
                        row_dict[f'col_{j}'] = val
                data.append(row_dict)
            else:
                data.append(row_data)
    
    return {
        'headers': headers,
        'rows': data,
        'row_count': len(data),
        'col_count': len(headers) if headers else 0
    }


def find_table_title(doc, table_index):
    """Try to find the title of a table by looking at nearby paragraphs."""
    # This is a simplified approach - in practice, finding table titles is complex
    # We'll try to extract from the first row if it looks like a title
    table = doc.tables[table_index]
    
    if len(table.rows) > 0:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        first_cell = first_row[0] if first_row else ""
        
        # Check if first cell contains "جدول" (table in Arabic)
        if 'جدول' in first_cell or 'الرقم' in first_cell:
            return first_cell[:100]
    
    return f"Table {table_index + 1}"


def match_table_to_item(table_title, table_data):
    """Try to match a table to an item based on its content."""
    title_lower = table_title.lower()
    
    # Known table patterns
    patterns = {
        '1-1': '1.1', 'جدول (1-1)': '1.1',
        '1-2': '1.5', 'جدول (1-2)': '1.5',
        '1-3': '1.9', 'جدول (1-3)': '1.9',
        '1-4': '1.9', 'جدول (1-4)': '1.9',
        '1-5': '1.9', 'جدول (1-5)': '1.9',
        '1-6': '1.10', 'جدول (1-6)': '1.10',
        '1-7': '1.10', 'جدول (1-7)': '1.10',
        '2-1': '2.2', 'جدول (2-1)': '2.2',
        '2-2': '2.5', 'جدول (2-2)': '2.5',
        '2-3': '2.5', 'جدول (2-3)': '2.5',
        '2-4': '2.5', 'جدول (2-4)': '2.5',
        '2-5': '2.5', 'جدول (2-5)': '2.5',
        '2-6': '2.7', 'جدول (2-6)': '2.7',
        '2-7': '2.7', 'جدول (2-7)': '2.7',
        '2-8': '2.7', 'جدول (2-8)': '2.7',
        '3-1': '3.1', 'جدول (3-1)': '3.1',
        '3-2': '3.1', 'جدول (3-2)': '3.1',
        '3-3': '3.1', 'جدول (3-3)': '3.1',
        '3-4': '3.3', 'جدول (3-4)': '3.3',
        '3-5': '3.3', 'جدول (3-5)': '3.3',
        '3-6': '3.3', 'جدول (3-6)': '3.3',
        '3-7': '3.3', 'جدول (3-7)': '3.3',
        '3-8': '3.7', 'جدول (3-8)': '3.7',
        '3-9': '3.7', 'جدول (3-9)': '3.7',
        '3-10': '3.7', 'جدول (3-10)': '3.7',
        '3-11': '3.7', 'جدول (3-11)': '3.7',
        '3-12': '3.7', 'جدول (3-12)': '3.7',
        '3-13': '3.8', 'جدول (3-13)': '3.8',
        '3-14': '3.8', 'جدول (3-14)': '3.8',
        '3-15': '3.8', 'جدول (3-15)': '3.8',
    }
    
    for pattern, item in patterns.items():
        if pattern in table_title:
            return item
    
    return None


def main():
    print("=" * 60)
    print("استخراج الجداول من التقرير الأصلي")
    print("=" * 60)
    
    if not REPORT_PATH.exists():
        print(f"ERROR: Report not found at {REPORT_PATH}")
        return
    
    print(f"Reading: {REPORT_PATH.name}")
    doc = Document(str(REPORT_PATH))
    
    print(f"Found {len(doc.tables)} tables")
    print()
    
    all_tables = {}
    item_tables = {}
    
    for i, table in enumerate(doc.tables):
        title = find_table_title(doc, i)
        data = extract_table_data(table)
        
        # Skip very small tables (probably formatting)
        if data['row_count'] < 2:
            continue
        
        # Skip index tables
        if 'فهرس' in title:
            continue
        
        table_info = {
            'index': i,
            'title': title,
            'rows': data['row_count'],
            'cols': data['col_count'],
            'headers': data['headers'],
            'data': data['rows'][:50]  # Limit to 50 rows
        }
        
        all_tables[f"table_{i}"] = table_info
        
        # Try to match to item
        item_code = match_table_to_item(title, data)
        if item_code:
            if item_code not in item_tables:
                item_tables[item_code] = []
            item_tables[item_code].append(table_info)
            print(f"✅ Table {i}: {data['row_count']}×{data['col_count']} → {item_code}")
        else:
            print(f"⚠️ Table {i}: {data['row_count']}×{data['col_count']} → ? ({title[:40]}...)")
    
    # Save all tables
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump({
            'total': len(all_tables),
            'matched': len(item_tables),
            'tables': all_tables,
            'by_item': item_tables
        }, f, ensure_ascii=False, indent=2)
    
    print()
    print(f"Saved to: {OUTPUT_FILE}")
    print(f"Total tables: {len(all_tables)}")
    print(f"Matched to items: {len(item_tables)}")
    
    # Save individual item tables
    for item_code, tables in item_tables.items():
        code_safe = item_code.replace('.', '_')
        item_dir = DATA_DIR / f"item_{code_safe}"
        item_dir.mkdir(exist_ok=True)
        
        tables_dir = item_dir / "tables"
        tables_dir.mkdir(exist_ok=True)
        
        for j, table in enumerate(tables):
            table_file = tables_dir / f"table_{j+1}.json"
            with open(table_file, 'w', encoding='utf-8') as f:
                json.dump(table, f, ensure_ascii=False, indent=2)
        
        print(f"  {item_code}: {len(tables)} tables saved")


if __name__ == '__main__':
    main()
