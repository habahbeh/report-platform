#!/usr/bin/env python3
"""
Export Item 3.1 skeleton to Word document.
Pilot script for the report generation system.
"""

import os
import sys
from pathlib import Path

# Add backend to path
backend_path = Path(__file__).parent.parent / 'backend'
sys.path.insert(0, str(backend_path))

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')

import django
django.setup()

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)


def set_arabic_font(run, font_name='Arial', size=12):
    """Set Arabic font for a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:cs'), font_name)


def create_rtl_paragraph(doc, text, font_size=12):
    """Create RTL paragraph with text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)
    run = p.add_run(text)
    set_arabic_font(run, 'Arial', font_size)
    return p


def export_skeleton_to_word(skeleton_path: str, output_path: str, images_dir: str):
    """Convert skeleton HTML to Word document."""
    
    # Read HTML
    with open(skeleton_path, 'r', encoding='utf-8') as f:
        html = f.read()
    
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()
    
    # Set RTL document
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_heading('البند 3.1: عدد الأبحاث المنشورة في SCOPUS و ISI', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in title.runs:
        set_arabic_font(run, 'Arial', 18)
    
    doc.add_paragraph()
    
    # Process components in order
    main = soup.find('main') or soup.find('body')
    if not main:
        print("Error: No main/body found in HTML")
        return
    
    components_processed = 0
    
    # Find all components
    for elem in main.find_all(['div', 'article']):
        data_component = elem.get('data-component')
        if not data_component:
            continue
        
        comp_type = data_component[0]  # p, c, t
        
        if comp_type == 'p':
            # Paragraph
            p_tags = elem.find_all('p')
            for p_tag in p_tags:
                text = p_tag.get_text().strip()
                if text:
                    create_rtl_paragraph(doc, text)
                    components_processed += 1
        
        elif comp_type == 'c':
            # Chart - add image
            img = elem.find('img')
            if img:
                src = img.get('src', '')
                img_path = os.path.join(images_dir, src)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    alt = img.get('alt', '')
                    if alt:
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cap.add_run(alt)
                        set_arabic_font(run, 'Arial', 10)
                        run.italic = True
                    
                    components_processed += 1
                    doc.add_paragraph()  # Spacer
                else:
                    print(f"Warning: Image not found: {img_path}")
        
        elif comp_type == 't':
            # Table
            table_title = elem.find(class_='table-title')
            html_table = elem.find('table')
            
            if table_title:
                p = create_rtl_paragraph(doc, table_title.get_text().strip(), 12)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
            
            if html_table:
                rows = html_table.find_all('tr')
                if rows:
                    # Count columns from first row
                    first_row = rows[0]
                    cols = first_row.find_all(['th', 'td'])
                    num_cols = len(cols)
                    
                    if num_cols > 0:
                        # Create Word table
                        word_table = doc.add_table(rows=len(rows), cols=num_cols)
                        word_table.style = 'Table Grid'
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['th', 'td'])
                            for j, cell in enumerate(cells):
                                if j < num_cols:
                                    word_cell = word_table.cell(i, j)
                                    word_cell.text = cell.get_text().strip()
                                    
                                    # Style cell
                                    for para in word_cell.paragraphs:
                                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        set_rtl_paragraph(para)
                                        for run in para.runs:
                                            set_arabic_font(run, 'Arial', 10)
                                            if i == 0:  # Header
                                                run.bold = True
                        
                        components_processed += 1
                        doc.add_paragraph()  # Spacer
    
    # Save
    doc.save(output_path)
    print(f"✅ Exported to: {output_path}")
    print(f"   Components processed: {components_processed}")
    
    return output_path


if __name__ == '__main__':
    data_dir = Path(__file__).parent.parent / 'data' / 'item_3_1'
    
    skeleton_path = data_dir / 'skeleton_3_1.html'
    output_path = data_dir / 'item_3_1_report.docx'
    
    if not skeleton_path.exists():
        print(f"Error: Skeleton not found: {skeleton_path}")
        sys.exit(1)
    
    export_skeleton_to_word(
        str(skeleton_path),
        str(output_path),
        str(data_dir)
    )
